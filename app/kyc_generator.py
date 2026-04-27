"""
Generates a styled KYC Word document from extracted document data.
NAAS — National Assurance & Advisory Services FZ LLC
17-section format per NAAS KYC Profile specification.
"""

import io
import re
import unicodedata
from datetime import date, datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Colour palette ─────────────────────────────────────────────────────────────
_NAVY       = "1B3A6B"
_BLUE       = "2E75B6"
_GREY_MED   = "555555"
_GREY_LITE  = "888888"
_WHITE      = "FFFFFF"
_OFF_WHITE  = "F5F5F5"
_DARK       = "1A1A1A"
_GREEN      = "1B6B3A"
_RED        = "C0392B"
_ORANGE     = "B45309"
_BORDER_CLR = "D0D7E0"

_TEXT_W = 16.0  # cm  (A4 with 2.5 cm side margins)


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _rgb(h: str) -> RGBColor:
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shd_cell(cell, fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:shd"))
    if old is not None:
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


def _cell_borders(cell, color: str = _BORDER_CLR, sz: int = 4, visible: bool = True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    borders = OxmlElement("w:tcBorders")
    val = "single" if visible else "none"
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        if visible:
            el.set(qn("w:sz"),    str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _get_tblPr(tbl):
    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tblPr


def _remove_tbl_borders(tbl):
    tblPr = _get_tblPr(tbl)
    old   = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _set_tbl_width(tbl, width_cm: float):
    tblPr = _get_tblPr(tbl)
    old   = tblPr.find(qn("w:tblW"))
    if old is not None:
        tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"),    str(int(width_cm * 567)))
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)


# ── Date helpers ──────────────────────────────────────────────────────────────

def _parse_date(s) -> date | None:
    if not s:
        return None
    for fmt in ("%Y-%m-%d", "%d %b %Y", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(str(s).strip(), fmt).date()
        except ValueError:
            continue
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})", str(s))
    if m:
        return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    return None


def _fmt_date(s) -> str:
    d = _parse_date(s)
    return d.strftime("%d %B %Y") if d else (str(s) if s else "—")


def _short_date(d: date | None) -> str:
    return d.strftime("%-d %b %Y") if d else "—"


def _expiry_label(expiry_str, today: date) -> tuple[str, str]:
    """Returns (formatted label, symbol ✓/⚠/✗)."""
    d = _parse_date(expiry_str)
    if d is None:
        return "—", "?"
    days = (d - today).days
    fmt  = d.strftime("%d %B %Y")
    if days < 0:
        mn = abs((today.year - d.year) * 12 + (today.month - d.month))
        return f"{fmt}  (EXPIRED — {mn} month{'s' if mn != 1 else ''} ago ✗)", "✗"
    elif days <= 30:
        return f"{fmt}  ({days} days remaining — EXPIRING SOON ⚠)", "⚠"
    else:
        return f"{fmt}  ({days} days remaining — VALID ✓)", "✓"


def _insurance_label(valid_to_str, today: date) -> tuple[str, str]:
    d = _parse_date(valid_to_str)
    if d is None:
        return "—", "?"
    days = (d - today).days
    fmt  = d.strftime("%d %B %Y")
    if days < 0:
        return f"{fmt}  (EXPIRED ✗)", "✗"
    elif days <= 30:
        return f"{fmt}  (EXPIRING SOON ⚠)", "⚠"
    else:
        return f"{fmt}  (ACTIVE ✓)", "✓"


# ── Value helpers ─────────────────────────────────────────────────────────────

def _s(val) -> str:
    """Safely stringify a value — handles lists the LLM may return instead of strings."""
    if val is None:
        return ""
    if isinstance(val, list):
        return " ".join(str(v) for v in val)
    return str(val)


def _v(d: dict, k: str, default: str = "—") -> str:
    val = d.get(k)
    return _s(val).strip() if val else default


def _normalise(s: str) -> str:
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]", "", s.lower())


def _name_tokens(s: str) -> set[str]:
    """Return a set of lowercased, ASCII-folded name parts (order-insensitive)."""
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode()
    return {t for t in re.sub(r"[^a-z\s]", "", s.lower()).split() if t}


def _names_match(a: str, b: str) -> bool:
    """Order-insensitive name comparison: same tokens = match."""
    na, nb = _normalise(a), _normalise(b)
    if na == nb or na in nb or nb in na:
        return True
    # Token-set comparison (handles reordered name parts)
    return _name_tokens(a) == _name_tokens(b)


def _match2(a: str, b: str) -> str:
    if not a or a == "—" or not b or b == "—":
        return "—"
    return "✓" if _names_match(a, b) else "⚠"


def _extract_percent(s: str) -> float | None:
    """Pull a percentage out of strings like '25.0%', '100%', '25 Shares AED 50,000 25%'.
    Returns the LAST percent found (most specific share %), or None."""
    if not s:
        return None
    matches = re.findall(r"(\d+(?:\.\d+)?)\s*%", str(s))
    if not matches:
        return None
    try:
        return float(matches[-1])
    except ValueError:
        return None


def _percent_match(a: str, b: str) -> str:
    """Compare two share-percentage strings numerically (tolerant of formatting)."""
    if not a or a == "—" or not b or b == "—":
        return "—"
    pa, pb = _extract_percent(a), _extract_percent(b)
    if pa is None or pb is None:
        # Fall back to name-style match if one side has no extractable %
        return _match2(a, b)
    return "✓" if abs(pa - pb) < 0.01 else "⚠"


def _match3(a: str, b: str, c: str) -> str:
    vals = [x for x in [a, b, c] if x and x != "—"]
    if not vals:
        return "—"
    if len(vals) == 1:
        return "✓"
    for i in range(len(vals)):
        for j in range(i + 1, len(vals)):
            if not _names_match(vals[i], vals[j]):
                return "⚠"
    return "✓"


# ── Partner identification ───────────────────────────────────────────────────

def identify_partners(extracted: dict) -> list[dict]:
    """
    Identify non-corporate partners from extracted data and check which
    already have matching personal documents (passport / EID / visa).

    Returns a list of partner dicts:
      [{"name", "nationality", "share_percentage", "has_passport",
        "has_emirates_id", "has_residence_visa"}, ...]

    Sources (in priority order): Partners Annex → MOA → Trade License.
    """
    partners: list[dict] = []

    # Source 1: Partners Annex
    pa = extracted.get("partners_annex") or {}
    pa_list = pa.get("partners", [])
    if isinstance(pa_list, list):
        for p in pa_list:
            if not isinstance(p, dict):
                continue
            if p.get("is_corporate"):
                continue
            partners.append({
                "name": _s(p.get("name", "")),
                "nationality": _s(p.get("nationality", "")),
                "share_percentage": _s(p.get("share_percentage", "")),
            })

    # Source 2: MOA shareholders[] array (or fallback to single owner_name)
    if not partners:
        moa = extracted.get("moa") or {}
        moa_shareholders = moa.get("shareholders") if isinstance(moa.get("shareholders"), list) else []
        moa_shareholders = [s for s in moa_shareholders
                            if isinstance(s, dict) and s.get("name")]
        if moa_shareholders:
            for sh in moa_shareholders:
                partners.append({
                    "name": _s(sh.get("name", "")),
                    "nationality": _s(sh.get("nationality", "")),
                    "share_percentage": _s(sh.get("share_percentage", "") or sh.get("shares", "")),
                })
        else:
            owner = moa.get("owner_name")
            if owner:
                partners.append({
                    "name": _s(owner),
                    "nationality": _s(moa.get("owner_nationality", "")),
                    "share_percentage": _s(moa.get("owner_shares", "")),
                })

    # Source 2b: MOA managers (if still no partners — multi-manager but no shareholder list)
    if not partners:
        moa = extracted.get("moa") or {}
        moa_managers = moa.get("managers") if isinstance(moa.get("managers"), list) else []
        for mgr in moa_managers:
            if isinstance(mgr, dict) and mgr.get("name"):
                partners.append({
                    "name": _s(mgr.get("name", "")),
                    "nationality": _s(mgr.get("nationality", "")),
                    "share_percentage": "",
                })

    # Source 3: Trade License owner
    if not partners:
        tl = extracted.get("trade_license") or {}
        owner = tl.get("owner_name")
        if owner:
            partners.append({
                "name": _s(owner),
                "nationality": _s(tl.get("owner_nationality", "")),
                "share_percentage": _s(tl.get("owner_share", "")),
            })

    if not partners:
        return []

    # Check which partners already have matching personal docs
    pp  = extracted.get("passport") or {}
    eid = extracted.get("emirates_id") or {}
    visa = extracted.get("residence_visa") or {}

    pp_name   = _s(pp.get("holder_name", ""))
    eid_name  = _s(eid.get("holder_name", ""))
    visa_name = _s(visa.get("holder_name", ""))

    for partner in partners:
        name = partner["name"]
        partner["has_passport"]       = bool(pp_name)   and _names_match(name, pp_name)
        partner["has_emirates_id"]    = bool(eid_name)  and _names_match(name, eid_name)
        partner["has_residence_visa"] = bool(visa_name) and _names_match(name, visa_name)

    return partners


# ── Document-building primitives ──────────────────────────────────────────────

def _spacer(doc, pts: float = 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    p._p.get_or_add_pPr().set(qn("w:jc"), "left")


def _section_header(doc, number: int, title: str):
    _spacer(doc, 10)
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)
    _set_tbl_width(tbl, _TEXT_W)
    cell = tbl.cell(0, 0)
    _shd_cell(cell, _NAVY)
    _cell_borders(cell, visible=False)
    cell.width = Cm(_TEXT_W)
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(5)
    para.paragraph_format.space_after  = Pt(5)
    para.paragraph_format.left_indent  = Pt(8)
    run = para.add_run(f"  {number}.   {title.upper()}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _rgb(_WHITE)
    _spacer(doc, 2)


def _sub_header(doc, title: str):
    _spacer(doc, 6)
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)
    _set_tbl_width(tbl, _TEXT_W)
    cell = tbl.cell(0, 0)
    _shd_cell(cell, "2E4F7A")
    _cell_borders(cell, visible=False)
    cell.width = Cm(_TEXT_W)
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Pt(10)
    run = para.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _rgb(_WHITE)
    _spacer(doc, 2)


def _kv_table(doc, rows: list[tuple[str, str]]):
    if not rows:
        return
    label_w = _TEXT_W * 0.36
    value_w = _TEXT_W * 0.64
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)
    _set_tbl_width(tbl, _TEXT_W)
    for i, (key, val) in enumerate(rows):
        fill = _WHITE if i % 2 == 0 else _OFF_WHITE
        row_el = tbl.rows[i]

        lc = row_el.cells[0]
        lc.width = Cm(label_w)
        _shd_cell(lc, fill)
        _cell_borders(lc, _BORDER_CLR)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after  = Pt(3)
        lp.paragraph_format.left_indent  = Pt(6)
        lr = lp.add_run(str(key) if key else "")
        lr.bold = True
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = _rgb(_NAVY)

        vc = row_el.cells[1]
        vc.width = Cm(value_w)
        _shd_cell(vc, fill)
        _cell_borders(vc, _BORDER_CLR)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after  = Pt(3)
        vp.paragraph_format.left_indent  = Pt(6)
        vr = vp.add_run(str(val) if val else "—")
        vr.font.size = Pt(9.5)
        vr.font.color.rgb = _rgb(_DARK)
    _spacer(doc, 6)


def _styled_cell(cell, txt: str, fill: str, w: float,
                 bold: bool = False, center: bool = False,
                 color: str = _DARK, pt: float = 9.5):
    cell.width = Cm(w)
    _shd_cell(cell, fill)
    _cell_borders(cell, _BORDER_CLR)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(6)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(txt) if txt else "—")
    r.font.size = Pt(pt)
    r.bold = bold
    r.font.color.rgb = _rgb(color)


def _match_color(sym: str) -> str:
    if "✓" in sym:
        return _GREEN
    if "⚠" in sym:
        return _ORANGE
    if "✗" in sym:
        return _RED
    return _GREY_MED


def _verify_table(doc, headers: tuple, rows: list[tuple]):
    """4-column cross-verification table."""
    if not rows:
        return
    col_w = [4.0, 4.5, 4.5, 3.0]
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)
    _set_tbl_width(tbl, _TEXT_W)
    for ci, (lbl, w) in enumerate(zip(headers, col_w)):
        cell = tbl.rows[0].cells[ci]
        _styled_cell(cell, lbl, _BLUE, w, bold=True, center=(ci == 3),
                     color=_WHITE, pt=9)
    for ri, (f, v1, v2, match) in enumerate(rows):
        fill = _WHITE if ri % 2 == 0 else _OFF_WHITE
        row  = tbl.rows[ri + 1]
        _styled_cell(row.cells[0], f,     fill, col_w[0], bold=True, color=_NAVY)
        _styled_cell(row.cells[1], v1,    fill, col_w[1], color=_DARK)
        _styled_cell(row.cells[2], v2,    fill, col_w[2], color=_DARK)
        mc = row.cells[3]
        _styled_cell(mc, match, fill, col_w[3], bold=("—" not in match),
                     center=True, color=_match_color(match), pt=11)
    _spacer(doc, 6)


def _three_way_table(doc, headers: tuple, rows: list[tuple]):
    """5-column table for 3-way cross-verification (Field | A | B | C | Match)."""
    if not rows:
        return
    col_w = [3.0, 3.2, 3.2, 3.2, 3.4]
    tbl = doc.add_table(rows=1 + len(rows), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)
    _set_tbl_width(tbl, _TEXT_W)
    for ci, (lbl, w) in enumerate(zip(headers, col_w)):
        _styled_cell(tbl.rows[0].cells[ci], lbl, _BLUE, w,
                     bold=True, center=(ci == 4), color=_WHITE, pt=9)
    for ri, row_data in enumerate(rows):
        f, v1, v2, v3, match = row_data
        fill = _WHITE if ri % 2 == 0 else _OFF_WHITE
        row  = tbl.rows[ri + 1]
        _styled_cell(row.cells[0], f,     fill, col_w[0], bold=True, color=_NAVY)
        _styled_cell(row.cells[1], v1,    fill, col_w[1], color=_DARK)
        _styled_cell(row.cells[2], v2,    fill, col_w[2], color=_DARK)
        _styled_cell(row.cells[3], v3,    fill, col_w[3], color=_DARK)
        _styled_cell(row.cells[4], match, fill, col_w[4], bold=("—" not in match),
                     center=True, color=_match_color(match), pt=11)
    _spacer(doc, 6)


def _doc_status_table(doc, rows: list[tuple]):
    """5-column personal document validity table.
    Columns: Document | Holder Name | Document No. | Expiry Date | Status"""
    if not rows:
        return
    headers = ("Document", "Holder Name", "Document No.", "Expiry Date", "Status")
    col_w   = [3.0, 3.8, 3.0, 3.2, 3.0]
    tbl = doc.add_table(rows=1 + len(rows), cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_tbl_borders(tbl)
    _set_tbl_width(tbl, _TEXT_W)
    for ci, (lbl, w) in enumerate(zip(headers, col_w)):
        _styled_cell(tbl.rows[0].cells[ci], lbl, _BLUE, w,
                     bold=True, center=(ci == 4), color=_WHITE, pt=9)
    for ri, (doc_type, holder, doc_no, expiry, sym) in enumerate(rows):
        fill = _WHITE if ri % 2 == 0 else _OFF_WHITE
        row  = tbl.rows[ri + 1]
        _styled_cell(row.cells[0], doc_type, fill, col_w[0], bold=True, color=_NAVY)
        _styled_cell(row.cells[1], holder,   fill, col_w[1], color=_DARK)
        _styled_cell(row.cells[2], doc_no,   fill, col_w[2], color=_DARK)
        _styled_cell(row.cells[3], expiry,   fill, col_w[3], color=_DARK)
        _styled_cell(row.cells[4], sym,      fill, col_w[4], bold=True,
                     center=True, color=_match_color(sym), pt=13)
    _spacer(doc, 6)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_kyc_document(extracted: dict, analysis: dict | None, today: date) -> bytes:
    """
    Build a styled KYC Word document and return bytes.

    extracted keys: trade_license | ejari | moa | insurance |
                    passport | emirates_id | residence_visa | vat_certificate |
                    board_resolution | poa | partners_annex |
                    certificate_of_incorporation | register_of_shareholders |
                    register_of_directors | certificate_of_good_standing

    `analysis` is the NAAS v4.0 compliance dict from `app.kyc_compliance.analyse`.
    Sections 14, 18, 19 read from it. Pass None for legacy callers — those
    sections then degrade to a "Not provided" line.
    """
    analysis = analysis or {}
    # ── Unpack ─────────────────────────────────────────────────────────────────
    tl_up   = bool(extracted.get("trade_license"))
    ej_up   = bool(extracted.get("ejari"))
    moa_up  = bool(extracted.get("moa"))
    ins_up  = bool(extracted.get("insurance"))
    pp_up   = bool(extracted.get("passport"))
    eid_up  = bool(extracted.get("emirates_id"))
    visa_up = bool(extracted.get("residence_visa"))
    vat_up  = bool(extracted.get("vat_certificate"))
    br_up   = bool(extracted.get("board_resolution"))
    poa_up  = bool(extracted.get("poa"))
    pa_up   = bool(extracted.get("partners_annex"))

    tl   = dict(extracted.get("trade_license")   or {})
    ej   = dict(extracted.get("ejari")            or {})
    moa  = dict(extracted.get("moa")              or {})
    ins  = dict(extracted.get("insurance")        or {})
    pp   = dict(extracted.get("passport")         or {})
    eid  = dict(extracted.get("emirates_id")      or {})
    visa = dict(extracted.get("residence_visa")   or {})
    vat  = dict(extracted.get("vat_certificate")  or {})
    br   = dict(extracted.get("board_resolution") or {})
    poa  = dict(extracted.get("poa")              or {})
    pa   = dict(extracted.get("partners_annex")   or {})

    for d in [tl, ej, moa, ins, pp, eid, visa, vat, br, poa, pa]:
        d.pop("error", None)

    # ── Key identifiers ────────────────────────────────────────────────────────
    company_name = (_v(tl, "company_name", "") or _v(moa, "company_name", "")
                    or _v(vat, "company_name", ""))
    legal_form   = _v(tl, "legal_form", "") or _v(moa, "legal_form", "")
    license_no   = _v(tl,  "license_number",  "—")
    ejari_no     = _v(ej,  "ejari_number",    "—")
    moa_no       = _v(moa, "contract_number", "—")

    # ── Expiry labels ──────────────────────────────────────────────────────────
    tl_label,   tl_sym   = _expiry_label(tl.get("expiry_date"),   today)
    ej_label,   ej_sym   = _expiry_label(ej.get("expiry_date"),   today)
    pp_label,   pp_sym   = _expiry_label(pp.get("expiry_date"),   today)
    eid_label,  eid_sym  = _expiry_label(eid.get("expiry_date"),  today)
    visa_label, visa_sym = _expiry_label(visa.get("expiry_date"), today)
    ins_label,  ins_sym  = _insurance_label(ins.get("valid_to"),  today)

    # ── Multi-partner personal docs ───────────────────────────────────────────
    partner_docs_list = extracted.get("partner_personal_docs") or []
    has_multi_partner = bool(partner_docs_list)

    # ── Flags accumulator ──────────────────────────────────────────────────────
    flags: list[dict] = []

    def _flag(ftype: str, docs: str, field: str, val_a: str, val_b: str, action: str):
        flags.append({"type": ftype, "docs": docs, "field": field,
                      "val_a": val_a, "val_b": val_b, "action": action})

    # ── Create document ────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # ── Page header & footer (NAAS spec) ───────────────────────────────────────
    version = (analysis.get("version") if isinstance(analysis, dict) else None) or "v1"
    today_short = today.strftime("%d %b %Y")
    header_text = (f"KYC PROFILE — {(company_name or '').upper()} | "
                   f"CONFIDENTIAL | {version} — {today_short}")
    footer_text = (f"Prepared by NAAS — National Assurance & Advisory Services FZ LLC | "
                   f"{today_short} | CONFIDENTIAL")

    hdr_para = sec.header.paragraphs[0]
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hdr_para.add_run(header_text)
    hr.font.size = Pt(8)
    hr.font.color.rgb = _rgb(_NAVY)
    hr.bold = True

    ftr_para = sec.footer.paragraphs[0]
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = ftr_para.add_run(footer_text)
    fr.font.size = Pt(8)
    fr.font.color.rgb = _rgb(_GREY_MED)

    def _centered(text, bold=False, pt=10, color=_DARK, sb=0, sa=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(pt)
        r.font.color.rgb = _rgb(color)

    # ── Title block ────────────────────────────────────────────────────────────
    _centered("KNOW YOUR CUSTOMER (KYC) PROFILE",
              bold=True, pt=16, color=_NAVY, sb=0, sa=4)
    _centered("National Assurance & Advisory Services FZ LLC (NAAS)",
              pt=9, color=_GREY_MED, sb=0, sa=3)
    _centered(company_name or "—", bold=True, pt=14, color=_BLUE, sb=0, sa=4)
    subtitle = f"{legal_form}  |  Dubai, UAE" if legal_form else "Dubai, UAE"
    _centered(subtitle, pt=10, color=_GREY_MED, sb=0, sa=4)
    ref_parts = []
    if moa_no   != "—": ref_parts.append(f"MOA: {moa_no}")
    if license_no != "—": ref_parts.append(f"Trade Licence: {license_no}")
    if ejari_no != "—": ref_parts.append(f"EJARI: {ejari_no}")
    ref_parts.append(f"Prepared: {today.strftime('%d %B %Y')}")
    _centered("  |  ".join(ref_parts), pt=9, color=_GREY_LITE, sb=0, sa=0)

    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(8)
    div.paragraph_format.space_after  = Pt(10)
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), _NAVY)
    pBdr.append(bot)
    div._p.get_or_add_pPr().append(pBdr)

    n = 1  # auto section counter

    # ══════════════════════════════════════════════════════════════════════════
    # 1 — COMPANY DETAILS
    # ══════════════════════════════════════════════════════════════════════════
    if tl or moa or vat:
        _section_header(doc, n, "Company Details"); n += 1
        rows: list[tuple] = []
        cn_en = (_v(tl, "company_name", "") or _v(moa, "company_name", "")
                 or _v(vat, "company_name", ""))
        cn_ar = (_v(tl, "company_name_arabic", "") or _v(moa, "company_name_arabic", "")
                 or _v(vat, "company_name_arabic", ""))
        rows.append(("Company Name (English)", cn_en or "—"))
        if cn_ar and cn_ar != "—":
            rows.append(("Company Name (Arabic)", cn_ar))
        lf = _v(tl, "legal_form", "") or _v(moa, "legal_form", "")
        if lf and lf != "—":
            rows.append(("Legal Type", lf))
        ia = _v(tl, "issuing_authority", "")
        if ia and ia != "—":
            rows.append(("Issuing Authority", ia))
        if moa.get("contract_number"):
            rows.append(("MOA Contract No.", _v(moa, "contract_number")))
        if moa.get("moa_date"):
            rows.append(("MOA Date", _fmt_date(moa["moa_date"])))
        if moa.get("company_duration"):
            rows.append(("Company Duration", _v(moa, "company_duration")))
        if moa.get("financial_year"):
            rows.append(("Financial Year", _v(moa, "financial_year")))
        if moa.get("disputes_jurisdiction"):
            rows.append(("Disputes Jurisdiction", _v(moa, "disputes_jurisdiction")))
        if vat.get("trn"):
            rows.append(("Tax Registration No. (TRN)", _v(vat, "trn")))
        _kv_table(doc, rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 2 — TRADE LICENCE DETAILS
    # ══════════════════════════════════════════════════════════════════════════
    if tl_up:
        _section_header(doc, n, "Trade Licence Details"); n += 1
        rows = [("Licence No.", _v(tl, "license_number"))]
        if tl.get("register_number"):
            rows.append(("Commercial Register No.", _v(tl, "register_number")))
        if tl.get("dcci_membership_number"):
            rows.append(("DCCI Membership No.", _v(tl, "dcci_membership_number")))
        if tl.get("license_type"):
            rows.append(("Licence Type", _v(tl, "license_type")))
        if tl.get("licence_category"):
            rows.append(("Licence Category", _v(tl, "licence_category")))
        if tl.get("issue_date"):
            rows.append(("Issue Date", _fmt_date(tl["issue_date"])))
        rows.append(("Expiry Date", tl_label))
        if tl.get("last_renewal_date"):
            rows.append(("Last Renewal Date", _fmt_date(tl["last_renewal_date"])))
        if tl.get("last_renewal_fee"):
            rows.append(("Last Renewal Fee", _v(tl, "last_renewal_fee")))
        if tl_sym in ("⚠", "✗"):
            _flag("Trade Licence Validity", "Trade Licence", "Expiry Date", tl_label, "",
                  "Renew Trade Licence immediately." if tl_sym == "✗"
                  else "Initiate renewal — expires within 30 days.")
        _kv_table(doc, rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 3 — REGISTERED ADDRESS & CONTACT DETAILS
    # ══════════════════════════════════════════════════════════════════════════
    addr_rows: list[tuple] = []
    if tl.get("registered_address"):
        addr_rows.append(("Registered Address", _v(tl, "registered_address")))
    if tl.get("unit_number"):
        addr_rows.append(("Unit No.", _v(tl, "unit_number")))
    if tl.get("building_name"):
        addr_rows.append(("Building", _v(tl, "building_name")))
    if tl.get("area"):
        addr_rows.append(("Area", _v(tl, "area")))
    if tl.get("parcel_id"):
        addr_rows.append(("Parcel ID / Land DM No.", _v(tl, "parcel_id")))
    if tl.get("makani_number"):
        addr_rows.append(("Makani No.", _v(tl, "makani_number")))
    if tl.get("phone_fax"):
        addr_rows.append(("Phone / Fax", _v(tl, "phone_fax")))
    if tl.get("mobile"):
        addr_rows.append(("Mobile No.", _v(tl, "mobile")))
    if tl.get("email"):
        addr_rows.append(("Email", _v(tl, "email")))
    if addr_rows:
        _section_header(doc, n, "Registered Address & Contact Details"); n += 1
        _kv_table(doc, addr_rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 4 — EJARI — TENANCY CONTRACT
    # ══════════════════════════════════════════════════════════════════════════
    if ej_up:
        _section_header(doc, n, "EJARI — Tenancy Contract"); n += 1
        rows = [("EJARI Contract No.", _v(ej, "ejari_number"))]
        if ej.get("registration_date"):
            rows.append(("Registration Date",           _fmt_date(ej["registration_date"])))
        if ej.get("registered_by"):
            rows.append(("Registered by",               _v(ej, "registered_by")))
        if ej.get("tenant_name"):
            rows.append(("Tenant Name",                 _v(ej, "tenant_name")))
        if ej.get("licence_number"):
            rows.append(("Trade Licence No. (EJARI)",   _v(ej, "licence_number")))
        if ej.get("licence_issuer"):
            rows.append(("Licence Issuer (EJARI)",      _v(ej, "licence_issuer")))
        if ej.get("start_date"):
            rows.append(("Lease Start Date",            _fmt_date(ej["start_date"])))
        rows.append(("Lease End Date",                  ej_label))
        if ej.get("annual_rent"):
            rows.append(("Annual Rent",                 _v(ej, "annual_rent")))
        if ej.get("security_deposit"):
            rows.append(("Security Deposit",            _v(ej, "security_deposit")))
        if ej.get("ejari_fees_paid"):
            rows.append(("EJARI Fees Paid",             _v(ej, "ejari_fees_paid")))
        if ej.get("unit_number"):
            rows.append(("Unit No.",                    _v(ej, "unit_number")))
        if ej.get("building_name"):
            rows.append(("Building",                    _v(ej, "building_name")))
        if ej.get("area"):
            rows.append(("Area",                        _v(ej, "area")))
        if ej.get("unit_type"):
            rows.append(("Unit Type",                   _v(ej, "unit_type")))
        if ej.get("size"):
            rows.append(("Size",                        _v(ej, "size")))
        if ej.get("plot_number"):
            rows.append(("Plot No.",                    _v(ej, "plot_number")))
        if ej.get("land_dm_parcel_id"):
            rows.append(("Land DM No. (Parcel ID)",     _v(ej, "land_dm_parcel_id")))
        if ej.get("makani_number"):
            rows.append(("Makani No.",                  _v(ej, "makani_number")))
        if ej_sym in ("⚠", "✗"):
            _flag("EJARI Validity", "EJARI Tenancy Contract", "Lease End Date", ej_label, "",
                  "Renew EJARI / tenancy contract." if ej_sym == "✗"
                  else "EJARI renewal due within 30 days.")
        _kv_table(doc, rows)

        landlord_rows: list[tuple] = []
        if ej.get("landlord_name"):
            landlord_rows.append(("Owner Name",           _v(ej, "landlord_name")))
        if ej.get("landlord_owner_number"):
            landlord_rows.append(("Owner No.",            _v(ej, "landlord_owner_number")))
        if ej.get("landlord_nationality"):
            landlord_rows.append(("Nationality",          _v(ej, "landlord_nationality")))
        if ej.get("property_manager"):
            landlord_rows.append(("Property Manager",     _v(ej, "property_manager")))
        if ej.get("property_manager_email"):
            landlord_rows.append(("Property Manager Email", _v(ej, "property_manager_email")))
        if landlord_rows:
            _sub_header(doc, "Landlord Details")
            _kv_table(doc, landlord_rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 5 — VAT REGISTRATION
    # ══════════════════════════════════════════════════════════════════════════
    if vat_up:
        _section_header(doc, n, "VAT Registration"); n += 1
        vat_rows: list[tuple] = []
        if vat.get("trn"):
            vat_rows.append(("Tax Registration No. (TRN)", _v(vat, "trn")))
        if vat.get("company_name"):
            vat_rows.append(("Registered Name",            _v(vat, "company_name")))
        if vat.get("company_name_arabic"):
            vat_rows.append(("Arabic Name",                _v(vat, "company_name_arabic")))
        if vat.get("effective_date"):
            vat_rows.append(("Effective Date",             _fmt_date(vat["effective_date"])))
        if vat.get("registered_address"):
            vat_rows.append(("VAT Registered Address",     _v(vat, "registered_address")))
        vat_rows.append(("Expiry",                         "No expiry date — ongoing registration"))
        if vat.get("return_period"):
            vat_rows.append(("Return Period",              _v(vat, "return_period")))
        if vat.get("registration_type"):
            vat_rows.append(("Registration Type",          _v(vat, "registration_type")))
        _kv_table(doc, vat_rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 6 — INSURANCE
    # ══════════════════════════════════════════════════════════════════════════
    if ins_up:
        _section_header(doc, n, "Insurance"); n += 1
        rows = []
        if ins.get("insurer"):
            rows.append(("Insurer",        _v(ins, "insurer")))
        if ins.get("policy_number"):
            rows.append(("Policy No.",     _v(ins, "policy_number")))
        if ins.get("insured_name"):
            rows.append(("Insured Name",   _v(ins, "insured_name")))
        if ins.get("coverage_type"):
            rows.append(("Coverage Type",  _v(ins, "coverage_type")))
        if ins.get("sum_insured"):
            rows.append(("Sum Insured",    _v(ins, "sum_insured")))
        if ins.get("premium"):
            rows.append(("Premium",        _v(ins, "premium")))
        if ins.get("deductible"):
            rows.append(("Deductible / Excess", _v(ins, "deductible")))
        if ins.get("valid_from"):
            rows.append(("Valid From",     _fmt_date(ins["valid_from"])))
        if ins.get("valid_to"):
            rows.append(("Valid To",       ins_label))
        if ins_sym in ("⚠", "✗"):
            _flag("Insurance Validity", "Insurance Certificate", "Valid To", ins_label, "",
                  "Renew insurance policy." if ins_sym == "✗"
                  else "Insurance expiring within 30 days.")
        _kv_table(doc, rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 7 — BUSINESS ACTIVITIES
    # ══════════════════════════════════════════════════════════════════════════
    if tl_up and tl.get("business_activity"):
        _section_header(doc, n, "Business Activities"); n += 1
        biz: list[tuple] = [("Primary Activity", _v(tl, "business_activity"))]
        if tl.get("activity_status"):
            biz.append(("Status",              _v(tl, "activity_status")))
        if tl.get("activity_scope"):
            biz.append(("Activity Scope",      _v(tl, "activity_scope")))
        if tl.get("regulatory_approval"):
            biz.append(("Regulatory Approval", _v(tl, "regulatory_approval")))
        _kv_table(doc, biz)

    # ══════════════════════════════════════════════════════════════════════════
    # 8 — SHARE CAPITAL & OWNERSHIP
    # ══════════════════════════════════════════════════════════════════════════
    if moa_up:
        _section_header(doc, n, "Share Capital & Ownership"); n += 1
        cap: list[tuple] = []
        if moa.get("share_capital"):
            cap.append(("Authorised & Paid-Up Capital", _v(moa, "share_capital")))
        if moa.get("shares_count"):
            cap.append(("Number of Shares",             _v(moa, "shares_count")))
        if moa.get("capital_currency"):
            cap.append(("Currency",                     _v(moa, "capital_currency")))
        if moa.get("capital_deposited"):
            cap.append(("Capital Deposited",            _v(moa, "capital_deposited")))
        if moa.get("statutory_reserve"):
            cap.append(("Statutory Reserve",            _v(moa, "statutory_reserve")))

        moa_shareholders = moa.get("shareholders") if isinstance(moa.get("shareholders"), list) else []
        moa_shareholders = [s for s in moa_shareholders if isinstance(s, dict) and s.get("name")]

        if len(moa_shareholders) > 1:
            for idx, sh in enumerate(moa_shareholders, 1):
                pct = _s(sh.get("share_percentage", "")).strip() or _s(sh.get("shares", "")).strip()
                label = f"Shareholder {idx}" + (f" ({pct})" if pct else "")
                cap.append((label, _s(sh.get("name", "")).strip() or "—"))
                if sh.get("shares") and sh.get("shares") != sh.get("share_percentage"):
                    cap.append((f"  Shareholding Detail {idx}", _s(sh.get("shares")).strip()))
        elif len(moa_shareholders) == 1:
            sh = moa_shareholders[0]
            pct = _s(sh.get("share_percentage", "")).strip() or _s(sh.get("shares", "")).strip()
            label = "Shareholder" + (f" ({pct})" if pct else "")
            cap.append((label, _s(sh.get("name", "")).strip() or "—"))
            if sh.get("shares"):
                cap.append(("Shareholding Detail", _s(sh.get("shares")).strip()))
        else:
            owner_nm = _v(moa, "owner_name", "") or _v(tl, "owner_name", "")
            if owner_nm and owner_nm != "—":
                pct = _s(moa.get("owner_shares", "")).strip()
                label = "Shareholder" + (f" ({pct})" if pct else "")
                cap.append((label, owner_nm))
            if moa.get("owner_shares"):
                cap.append(("Shareholding Detail", _v(moa, "owner_shares")))
        _kv_table(doc, cap)

    # ══════════════════════════════════════════════════════════════════════════
    # 9 — OWNER / SHAREHOLDER DETAILS
    # ══════════════════════════════════════════════════════════════════════════
    if tl_up or moa_up or pp_up or eid_up:
        _section_header(doc, n, "Owner / Shareholder Details"); n += 1
        moa_shareholders = moa.get("shareholders") if isinstance(moa.get("shareholders"), list) else []
        moa_shareholders = [s for s in moa_shareholders if isinstance(s, dict) and s.get("name")]
        rows: list[tuple] = []

        if len(moa_shareholders) > 1:
            for idx, sh in enumerate(moa_shareholders, 1):
                if idx > 1:
                    rows.append(("", ""))
                    rows.append((f"─── SHAREHOLDER {idx} ─────────────────────────────", ""))
                else:
                    rows.append((f"─── SHAREHOLDER {idx} ─────────────────────────────", ""))
                rows.append(("Full Name (English)", _s(sh.get("name", "")).strip() or "—"))
                if sh.get("name_arabic"):
                    rows.append(("Full Name (Arabic)", _s(sh["name_arabic"]).strip()))
                if sh.get("nationality"):
                    rows.append(("Nationality", _s(sh["nationality"]).strip()))
                if sh.get("person_number"):
                    rows.append(("Person No.", _s(sh["person_number"]).strip()))
                if sh.get("shares"):
                    rows.append(("Shareholding", _s(sh["shares"]).strip()))
                elif sh.get("share_percentage"):
                    rows.append(("Shareholding", _s(sh["share_percentage"]).strip()))
                if sh.get("liability"):
                    rows.append(("Liability", _s(sh["liability"]).strip()))
                if sh.get("residence"):
                    rows.append(("Residence", _s(sh["residence"]).strip()))
        else:
            owner_name = (
                _s(moa.get("owner_name")).strip() or
                _s(tl.get("owner_name")).strip() or
                _s(pp.get("holder_name")).strip() or
                _s(eid.get("holder_name")).strip()
            )
            rows.append(("Full Name (English)", owner_name or "—"))
            if moa.get("owner_name_arabic"):
                rows.append(("Full Name (Arabic)",   _v(moa, "owner_name_arabic")))
            nat = (_v(moa, "owner_nationality", "") or _v(tl, "owner_nationality", "")
                   or _v(pp, "nationality", "") or _v(eid, "nationality", ""))
            if nat and nat != "—":
                rows.append(("Nationality", nat))
            pno = _v(moa, "owner_person_number", "") or _v(tl, "owner_person_number", "")
            if pno and pno != "—":
                rows.append(("Person No. (Licence)", pno))
            if moa.get("owner_shares"):
                rows.append(("Shareholding",  _v(moa, "owner_shares")))
            if moa.get("owner_liability"):
                rows.append(("Liability",     _v(moa, "owner_liability")))
            if moa.get("owner_residence"):
                rows.append(("Residence",     _v(moa, "owner_residence")))

        if has_multi_partner:
            # ── Per-partner personal documents ────────────────────────────
            _kv_table(doc, rows)
            for pidx, pdoc in enumerate(partner_docs_list, 1):
                pname = pdoc.get("partner_name", f"Partner {pidx}")
                ppp   = dict(pdoc.get("passport") or {})
                peid  = dict(pdoc.get("emirates_id") or {})
                pvisa = dict(pdoc.get("residence_visa") or {})
                ppp.pop("error", None); peid.pop("error", None); pvisa.pop("error", None)
                if not (ppp or peid or pvisa):
                    continue
                _sub_header(doc, f"Personal Documents — {pname}")
                p_rows: list[tuple] = []
                if ppp:
                    p_rows.append(("─── PASSPORT ───────────────────────────────────────", ""))
                    if ppp.get("holder_name"):
                        p_rows.append(("Name", _v(ppp, "holder_name")))
                    if ppp.get("passport_number"):
                        p_rows.append(("Passport No.",   _v(ppp, "passport_number")))
                    if ppp.get("date_of_birth"):
                        p_rows.append(("Date of Birth",  _fmt_date(ppp["date_of_birth"])))
                    if ppp.get("place_of_birth"):
                        p_rows.append(("Place of Birth", _v(ppp, "place_of_birth")))
                    if ppp.get("issue_date"):
                        p_rows.append(("Issue Date",     _fmt_date(ppp["issue_date"])))
                    if ppp.get("expiry_date"):
                        ppp_lbl, _ = _expiry_label(ppp["expiry_date"], today)
                        p_rows.append(("Expiry Date",    ppp_lbl))
                if peid:
                    p_rows.append(("─── EMIRATES ID ────────────────────────────────────", ""))
                    if peid.get("holder_name"):
                        p_rows.append(("Name",          _v(peid, "holder_name")))
                    if peid.get("id_number"):
                        p_rows.append(("ID No.",        _v(peid, "id_number")))
                    if peid.get("date_of_birth"):
                        p_rows.append(("Date of Birth", _fmt_date(peid["date_of_birth"])))
                    if peid.get("expiry_date"):
                        peid_lbl, _ = _expiry_label(peid["expiry_date"], today)
                        p_rows.append(("Expiry Date",   peid_lbl))
                if pvisa:
                    p_rows.append(("─── UAE RESIDENCE VISA ─────────────────────────────", ""))
                    if pvisa.get("holder_name"):
                        p_rows.append(("Name",                 _v(pvisa, "holder_name")))
                    if pvisa.get("visa_number"):
                        p_rows.append(("Visa / Permit No.",    _v(pvisa, "visa_number")))
                    if pvisa.get("file_number"):
                        p_rows.append(("File No.",             _v(pvisa, "file_number")))
                    if pvisa.get("uid_number"):
                        p_rows.append(("Unified No. (UID)",    _v(pvisa, "uid_number")))
                    if pvisa.get("profession"):
                        p_rows.append(("Profession",           _v(pvisa, "profession")))
                    if pvisa.get("employer"):
                        p_rows.append(("Sponsor / Employer",   _v(pvisa, "employer")))
                    if pvisa.get("place_of_issue"):
                        p_rows.append(("Place of Issue",       _v(pvisa, "place_of_issue")))
                    if pvisa.get("issue_date"):
                        p_rows.append(("Issue Date",           _fmt_date(pvisa["issue_date"])))
                    if pvisa.get("expiry_date"):
                        pvisa_lbl, _ = _expiry_label(pvisa["expiry_date"], today)
                        p_rows.append(("Expiry Date",          pvisa_lbl))
                if p_rows:
                    _kv_table(doc, p_rows)
        else:
            # ── Single-partner personal documents (original logic) ────────
            if pp_up and pp:
                rows.append(("─── PASSPORT ───────────────────────────────────────", ""))
                if pp.get("passport_number"):
                    rows.append(("Passport No.",    _v(pp, "passport_number")))
                if pp.get("date_of_birth"):
                    rows.append(("Date of Birth",  _fmt_date(pp["date_of_birth"])))
                if pp.get("place_of_birth"):
                    rows.append(("Place of Birth", _v(pp, "place_of_birth")))
                if pp.get("issue_date"):
                    rows.append(("Issue Date",     _fmt_date(pp["issue_date"])))
                if pp.get("expiry_date"):
                    rows.append(("Expiry Date",    pp_label))

            if eid_up and eid:
                rows.append(("─── EMIRATES ID ────────────────────────────────────", ""))
                if eid.get("id_number"):
                    rows.append(("ID No.",        _v(eid, "id_number")))
                if eid.get("date_of_birth"):
                    rows.append(("Date of Birth", _fmt_date(eid["date_of_birth"])))
                if eid.get("expiry_date"):
                    rows.append(("Expiry Date",   eid_label))

            if visa_up and visa:
                rows.append(("─── UAE RESIDENCE VISA ─────────────────────────────", ""))
                if visa.get("visa_number"):
                    rows.append(("Visa / Permit No.",   _v(visa, "visa_number")))
                if visa.get("file_number"):
                    rows.append(("File No.",            _v(visa, "file_number")))
                if visa.get("uid_number"):
                    rows.append(("Unified No. (UID)",   _v(visa, "uid_number")))
                if visa.get("profession"):
                    rows.append(("Profession",          _v(visa, "profession")))
                if visa.get("employer"):
                    rows.append(("Sponsor / Employer",  _v(visa, "employer")))
                if visa.get("place_of_issue"):
                    rows.append(("Place of Issue",      _v(visa, "place_of_issue")))
                if visa.get("issue_date"):
                    rows.append(("Issue Date",          _fmt_date(visa["issue_date"])))
                if visa.get("expiry_date"):
                    rows.append(("Expiry Date",         visa_label))
            _kv_table(doc, rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 10 — MANAGEMENT DETAILS
    # ══════════════════════════════════════════════════════════════════════════
    if moa_up or (tl_up and tl.get("manager_name")):
        _section_header(doc, n, "Management Details"); n += 1
        moa_managers = moa.get("managers") if isinstance(moa.get("managers"), list) else []
        moa_managers = [m for m in moa_managers if isinstance(m, dict) and m.get("name")]
        rows = []

        if len(moa_managers) > 1:
            for idx, mgr in enumerate(moa_managers, 1):
                if idx > 1:
                    rows.append(("", ""))
                rows.append((f"─── MANAGER {idx} ─────────────────────────────────", ""))
                rows.append(("Manager Name (English)", _s(mgr.get("name", "")).strip() or "—"))
                if mgr.get("name_arabic"):
                    rows.append(("Manager Name (Arabic)", _s(mgr["name_arabic"]).strip()))
                if mgr.get("nationality"):
                    rows.append(("Nationality", _s(mgr["nationality"]).strip()))
                if mgr.get("person_number"):
                    rows.append(("Person No.", _s(mgr["person_number"]).strip()))
                if mgr.get("role"):
                    rows.append(("Role", _s(mgr["role"]).strip()))
                if mgr.get("appointment_term"):
                    rows.append(("Appointment Term", _s(mgr["appointment_term"]).strip()))
                if mgr.get("residence"):
                    rows.append(("Residence", _s(mgr["residence"]).strip()))
                if mgr.get("pobox"):
                    rows.append(("P.O. Box", _s(mgr["pobox"]).strip()))
            if moa.get("signing_authority"):
                rows.append(("", ""))
                rows.append(("Signing Authority", _v(moa, "signing_authority")))
        else:
            mgr_name = _v(moa, "manager_name", "") or _v(tl, "manager_name", "")
            rows.append(("Manager Name (English)", mgr_name or "—"))
            if moa.get("manager_name_arabic"):
                rows.append(("Manager Name (Arabic)", _v(moa, "manager_name_arabic")))
            mgr_nat = _v(moa, "manager_nationality", "") or _v(tl, "manager_nationality", "")
            if mgr_nat and mgr_nat != "—":
                rows.append(("Nationality", mgr_nat))
            mgr_pno = _v(moa, "manager_person_number", "") or _v(tl, "manager_person_number", "")
            if mgr_pno and mgr_pno != "—":
                rows.append(("Person No. (Licence)", mgr_pno))
            if moa.get("manager_role"):
                rows.append(("Role",             _v(moa, "manager_role")))
            if moa.get("manager_appointment_term"):
                rows.append(("Appointment Term", _v(moa, "manager_appointment_term")))
            if moa.get("manager_residence"):
                rows.append(("Residence",        _v(moa, "manager_residence")))
            if moa.get("manager_pobox"):
                rows.append(("P.O. Box",         _v(moa, "manager_pobox")))
            if moa.get("signing_authority"):
                rows.append(("Signing Authority", _v(moa, "signing_authority")))
        _kv_table(doc, rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 11 — BANKING & SIGNATORY AUTHORITY
    # ══════════════════════════════════════════════════════════════════════════
    if moa_up:
        _section_header(doc, n, "Banking & Signatory Authority"); n += 1
        _moa_mgrs_l = moa.get("managers") if isinstance(moa.get("managers"), list) else []
        _moa_mgrs_l = [m for m in _moa_mgrs_l if isinstance(m, dict) and m.get("name")]
        _all_mgr_names = ", ".join(_s(m.get("name")).strip() for m in _moa_mgrs_l)
        signatory = (_v(moa, "authorised_signatory", "")
                     or _all_mgr_names
                     or _v(moa, "manager_name", "")
                     or _v(tl,  "manager_name", ""))
        banking: list[tuple] = [
            ("Authorised Signatory", signatory or "—"),
            ("Signing Mode",         _v(moa, "signing_mode")),
        ]
        for key, label in [
            ("bank_open_close", "Open / Close Bank Accounts"),
            ("bank_operate",    "Operate Bank Accounts"),
            ("bank_cheques",    "Sign Cheques"),
            ("bank_transfer",   "Transfer / Withdraw Funds"),
            ("bank_tenders",    "Sign Tenders & Contracts"),
            ("bank_lc",         "Issue Letters of Credit"),
            ("bank_vat",        "VAT / FTA Returns"),
            ("bank_delegate",   "Delegate Authority"),
        ]:
            if moa.get(key):
                banking.append((label, _v(moa, key)))
        _kv_table(doc, banking)

    # ══════════════════════════════════════════════════════════════════════════
    # N — BOARD RESOLUTION STATUS
    # ══════════════════════════════════════════════════════════════════════════
    # Determine if MOA grants banking authority or Board Resolution is required
    moa_banking_sufficient = False
    br_provided = br_up and br
    if moa_up:
        _section_header(doc, n, "Board Resolution Status"); n += 1
        # Check if MOA explicitly grants banking authority
        has_bank_open   = bool(moa.get("bank_open_close"))
        has_bank_cheque = bool(moa.get("bank_cheques"))
        has_bank_xfer   = bool(moa.get("bank_transfer"))
        moa_banking_sufficient = has_bank_open or has_bank_cheque or has_bank_xfer

        br_rows: list[tuple] = []
        br_rows.append(("MOA Type", "Original MOA — no amendments detected"))
        _br_mgrs_l = moa.get("managers") if isinstance(moa.get("managers"), list) else []
        _br_mgrs_l = [m for m in _br_mgrs_l if isinstance(m, dict) and m.get("name")]
        _br_all_names = ", ".join(_s(m.get("name")).strip() for m in _br_mgrs_l)
        signatory = (_v(moa, "authorised_signatory", "")
                     or _br_all_names
                     or _v(moa, "manager_name", "")
                     or _v(tl,  "manager_name", ""))
        br_rows.append(("Authorised Signatory", signatory or "—"))
        br_rows.append(("Signing Mode", _v(moa, "signing_mode", "—")))
        br_rows.append(("Bank Account Opening",
                        f"✓ Authorised — per MOA" if has_bank_open
                        else "✗ Not stated in MOA"))
        br_rows.append(("Cheque Signing",
                        f"✓ Authorised — per MOA" if has_bank_cheque
                        else "✗ Not stated in MOA"))
        br_rows.append(("Fund Transfer",
                        f"✓ Authorised — per MOA" if has_bank_xfer
                        else "✗ Not stated in MOA"))
        br_rows.append(("Delegate via POA",
                        f"✓ Permitted — per MOA" if moa.get("bank_delegate")
                        else "✗ Not stated in MOA"))

        if moa_banking_sufficient:
            br_rows.append(("Board Resolution Required",
                            "✗ No — MOA is sufficient. Banking authority explicitly granted."))
        else:
            br_rows.append(("Board Resolution Required",
                            "✓ Yes — MOA does not explicitly grant banking authority."))
            if not br_provided:
                _flag("Banking Authority Missing",
                      f"MOA (Contract No. {_v(moa, 'contract_number', '—')})",
                      "Banking Authority",
                      "MOA does not explicitly grant banking/signatory authority to the Manager.",
                      "",
                      f"Provide notarised Board Resolution / Owner's Resolution authorising "
                      f"{signatory or 'the Manager'} to open, operate, and sign on company bank accounts.")

        if br_provided:
            br_rows.append(("", ""))  # spacer
            br_rows.append(("─── BOARD RESOLUTION PROVIDED ───────────────────", ""))
            br_rows.append(("Resolution Type", _v(br, "resolution_type")))
            br_rows.append(("Resolution Date", _fmt_date(br.get("resolution_date"))))
            br_rows.append(("Named Signatory", _v(br, "signatory_name")))
            br_rows.append(("Designation", _v(br, "signatory_designation")))
            br_rows.append(("Signing Mode", _v(br, "signing_mode")))
            br_rows.append(("Named Bank(s)", _v(br, "named_banks")))

            powers = []
            if br.get("bank_open_close"): powers.append("Open/Close Accounts")
            if br.get("bank_operate"):    powers.append("Operate Accounts")
            if br.get("bank_cheques"):    powers.append("Sign Cheques")
            if br.get("bank_transfer"):   powers.append("Transfer/Withdraw Funds")
            if br.get("bank_sign_documents"): powers.append("Sign Banking Documents")
            br_rows.append(("Powers Granted", ", ".join(powers) if powers else "—"))
            br_rows.append(("Notarised", "✓ Yes" if br.get("notarised") else "✗ No / Not confirmed"))
            br_rows.append(("Company Stamp", "✓ Yes" if br.get("company_stamp") else "✗ No / Not confirmed"))

            if br.get("validity_period"):
                br_rows.append(("Validity Period", _v(br, "validity_period")))
            if br.get("expiry_date"):
                br_exp_label, br_exp_sym = _expiry_label(br.get("expiry_date"), today)
                br_rows.append(("Expiry Date", br_exp_label))
                if br_exp_sym in ("⚠", "✗"):
                    _flag("Board Resolution Validity",
                          "Board Resolution", "Expiry Date", br_exp_label, "",
                          "Renew Board Resolution." if br_exp_sym == "✗"
                          else "Board Resolution expiring within 30 days.")

        _kv_table(doc, br_rows)
    elif br_provided:
        # Board resolution uploaded without MOA
        _section_header(doc, n, "Board Resolution"); n += 1
        br_rows = [
            ("Resolution Type", _v(br, "resolution_type")),
            ("Resolution Date", _fmt_date(br.get("resolution_date"))),
            ("Company Name", _v(br, "company_name")),
            ("Licence No.", _v(br, "licence_number")),
            ("Named Signatory", _v(br, "signatory_name")),
            ("Designation", _v(br, "signatory_designation")),
            ("Signing Mode", _v(br, "signing_mode")),
            ("Named Bank(s)", _v(br, "named_banks")),
            ("Notarised", "✓ Yes" if br.get("notarised") else "✗ No / Not confirmed"),
        ]
        if br.get("validity_period"):
            br_rows.append(("Validity Period", _v(br, "validity_period")))
        if br.get("expiry_date"):
            br_exp_label, br_exp_sym = _expiry_label(br.get("expiry_date"), today)
            br_rows.append(("Expiry Date", br_exp_label))
        _kv_table(doc, br_rows)

    # ══════════════════════════════════════════════════════════════════════════
    # N — PHYSICAL PRESENCE & POA STATUS
    # ══════════════════════════════════════════════════════════════════════════
    if eid_up or visa_up or pp_up or poa_up:
        _section_header(doc, n, "Physical Presence & POA Status"); n += 1

        signatory_name = (
            _v(moa, "authorised_signatory", "") or
            _v(moa, "manager_name", "") or _v(tl, "manager_name", "") or
            _v(br, "signatory_name", "") or
            _v(moa, "owner_name", "") or _v(tl, "owner_name", "") or
            _v(pp, "holder_name", "") or _v(eid, "holder_name", "")
        )
        authority_source = "MOA" if moa_banking_sufficient else (
            "Board Resolution" if br_provided else "Not confirmed"
        )

        eid_valid = eid_up and eid and eid_sym == "✓"
        visa_valid = visa_up and visa and visa_sym == "✓"
        pp_valid = pp_up and pp and pp_sym == "✓"
        all_valid = eid_valid and visa_valid and pp_valid
        in_uae = visa_valid  # Proxy: valid visa implies UAE residency

        can_proceed = all_valid and (moa_banking_sufficient or br_provided)

        presence_rows: list[tuple] = [
            ("Authorised Signatory", signatory_name or "—"),
            ("Authority Source", authority_source),
            ("In UAE (Resident)", "✓ Yes — valid UAE Residence Visa" if in_uae
             else ("⚠ Not confirmed — no valid Visa uploaded" if not visa_up
                   else "✗ Visa expired — cannot confirm UAE residency")),
            ("Emirates ID Valid", f"✓ Valid" if eid_valid
             else ("✗ Expired" if eid_up and eid_sym == "✗"
                   else ("⚠ Expiring Soon" if eid_up and eid_sym == "⚠"
                         else "— Not uploaded"))),
            ("Passport Valid", f"✓ Valid" if pp_valid
             else ("✗ Expired" if pp_up and pp_sym == "✗"
                   else ("⚠ Expiring Soon" if pp_up and pp_sym == "⚠"
                         else "— Not uploaded"))),
            ("UAE Visa Valid", f"✓ Valid" if visa_valid
             else ("✗ Expired" if visa_up and visa_sym == "✗"
                   else ("⚠ Expiring Soon" if visa_up and visa_sym == "⚠"
                         else "— Not uploaded"))),
        ]

        if can_proceed:
            presence_rows.append(("Can Proceed", "✓ Yes — signatory can attend bank in person"))
            presence_rows.append(("Action Required", "None — bring original EID + Passport + Visa + company documents"))
        elif not (moa_banking_sufficient or br_provided):
            presence_rows.append(("Can Proceed", "✗ No — banking authority not confirmed"))
            presence_rows.append(("Action Required", "Board Resolution required (see above)"))
            _flag("Signatory Authority Not Confirmed",
                  "MOA / Board Resolution", "Banking Authority",
                  "No banking authority confirmed via MOA or Board Resolution.", "",
                  "Provide Board Resolution or confirm MOA banking powers.")
        elif not all_valid:
            expired_docs = []
            if eid_up and eid_sym in ("⚠", "✗"): expired_docs.append("Emirates ID")
            if pp_up and pp_sym in ("⚠", "✗"):   expired_docs.append("Passport")
            if visa_up and visa_sym in ("⚠", "✗"): expired_docs.append("UAE Visa")
            missing_docs = []
            if not eid_up:  missing_docs.append("Emirates ID")
            if not pp_up:   missing_docs.append("Passport")
            if not visa_up: missing_docs.append("UAE Visa")
            issues = []
            if expired_docs: issues.append(f"Expired/expiring: {', '.join(expired_docs)}")
            if missing_docs: issues.append(f"Missing: {', '.join(missing_docs)}")
            presence_rows.append(("Can Proceed", f"⚠ Blocked — {'; '.join(issues)}"))
            presence_rows.append(("Action Required",
                                  "Renew expired documents before attending bank. "
                                  "Alternatively, execute notarised POA to a UAE-resident individual."))

        # POA section
        if poa_up and poa:
            presence_rows.append(("", ""))
            presence_rows.append(("─── POWER OF ATTORNEY ──────────────────────────", ""))
            presence_rows.append(("POA Status", "✓ POA Provided"))
            presence_rows.append(("Grantor", _v(poa, "grantor_name")))
            presence_rows.append(("Grantee (Attorney)", _v(poa, "grantee_name")))
            presence_rows.append(("Grantee Nationality", _v(poa, "grantee_nationality")))
            if poa.get("scope_description"):
                presence_rows.append(("Scope", _v(poa, "scope_description")))
            if poa.get("named_banks"):
                presence_rows.append(("Named Bank(s)", _v(poa, "named_banks")))
            if poa.get("poa_date"):
                presence_rows.append(("POA Date", _fmt_date(poa["poa_date"])))
            if poa.get("validity_period"):
                presence_rows.append(("Validity Period", _v(poa, "validity_period")))
            if poa.get("expiry_date"):
                poa_exp_label, poa_exp_sym = _expiry_label(poa.get("expiry_date"), today)
                presence_rows.append(("Expiry Date", poa_exp_label))
                if poa_exp_sym in ("⚠", "✗"):
                    _flag("POA Validity", "Power of Attorney", "Expiry Date",
                          poa_exp_label, "",
                          "POA expired — execute new POA." if poa_exp_sym == "✗"
                          else "POA expiring within 30 days.")
            presence_rows.append(("Notarised", "✓ Yes" if poa.get("notarised") else "✗ No / Not confirmed"))
            if poa.get("attestation_status"):
                presence_rows.append(("Attestation", _v(poa, "attestation_status")))
            presence_rows.append(("Language", _v(poa, "language")))
            presence_rows.append(("Governing Law", _v(poa, "governing_law")))

            # Check POA grantee documents
            _sub_header(doc, "POA Grantee Document Requirements")
            grantee_rows: list[tuple] = [
                ("Grantee Name", _v(poa, "grantee_name")),
                ("UAE Resident", "⚠ Must be verified — valid UAE Residence Visa required"),
                ("Valid Emirates ID", "⚠ Must be verified — copy required for KYC"),
                ("Valid Passport", "⚠ Must be verified — copy required for KYC"),
                ("Age 21+", "⚠ Must be verified"),
                ("Not Company Auditor", "⚠ Must be verified — no conflict of interest"),
            ]
            _kv_table(doc, grantee_rows)
            _flag("POA Grantee Documents Not Verified",
                  "Power of Attorney", "Grantee Documents",
                  f"POA provided but grantee ({_v(poa, 'grantee_name')}) personal documents "
                  "have not been submitted for verification.", "",
                  f"Provide valid EID, Passport, and UAE Residence Visa of "
                  f"{_v(poa, 'grantee_name')} for KYC file.")

        _kv_table(doc, presence_rows)

    # ══════════════════════════════════════════════════════════════════════════
    # N — PARTNERS / SHAREHOLDERS
    # ══════════════════════════════════════════════════════════════════════════
    if pa_up and pa:
        _section_header(doc, n, "Partners / Shareholders — Annex"); n += 1
        pa_rows: list[tuple] = []
        if pa.get("company_name"):
            pa_rows.append(("Company Name", _v(pa, "company_name")))
        if pa.get("licence_number"):
            pa_rows.append(("Licence No.", _v(pa, "licence_number")))

        partners = pa.get("partners", [])
        if isinstance(partners, list):
            corporate_partners = []
            for idx, partner in enumerate(partners, 1):
                if not isinstance(partner, dict):
                    continue
                pa_rows.append(("", ""))
                pa_rows.append((f"─── PARTNER {idx} ─────────────────────────────────", ""))
                pa_rows.append(("Name", _v(partner, "name")))
                if partner.get("name_arabic"):
                    pa_rows.append(("Name (Arabic)", _v(partner, "name_arabic")))
                pa_rows.append(("Nationality", _v(partner, "nationality")))
                if partner.get("person_number"):
                    pa_rows.append(("Person No.", _v(partner, "person_number")))
                pa_rows.append(("Shareholding", _v(partner, "share_percentage")))
                if partner.get("share_value"):
                    pa_rows.append(("Share Value", _v(partner, "share_value")))
                if partner.get("role"):
                    pa_rows.append(("Role", _v(partner, "role")))

                is_corp = partner.get("is_corporate", False)
                pa_rows.append(("Shareholder Type",
                                "CORPORATE ENTITY" if is_corp else "Natural Person"))
                if is_corp:
                    if partner.get("jurisdiction"):
                        pa_rows.append(("Jurisdiction", _v(partner, "jurisdiction")))
                    corporate_partners.append(partner)

            # Flag corporate shareholders
            for cp in corporate_partners:
                cp_name = _v(cp, "name")
                cp_share = _v(cp, "share_percentage", "")
                _flag("Corporate Shareholder Identified",
                      "Partners Annex", "Shareholder Type",
                      f"{cp_name} holds {cp_share} — enhanced KYC required.",
                      "",
                      f"Full corporate KYC required for {cp_name}: "
                      "Certificate of Incorporation, MOA/AOA, Register of Shareholders, "
                      "Register of Directors, Certificate of Good Standing, Board Resolution, "
                      "and UBO/Director passports. "
                      "All foreign documents must complete the 4-stage UAE attestation chain "
                      "(Translation → Home MFA → UAE Embassy → UAE MOFA).")
        _kv_table(doc, pa_rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 14 — CORPORATE SHAREHOLDER KYC (NAAS spec Step 7)
    # ══════════════════════════════════════════════════════════════════════════
    corp_kyc = analysis.get("corporate_kyc") or [] if isinstance(analysis, dict) else []
    if corp_kyc:
        _section_header(doc, n, "Corporate Shareholder KYC"); n += 1
        _country_notes = {
            "uk": "UK: Apostille → UAE Embassy London → UAE MOFA.",
            "united kingdom": "UK: Apostille → UAE Embassy London → UAE MOFA.",
            "usa": "USA: State Notary → State MFA → UAE Embassy → UAE MOFA.",
            "united states": "USA: State Notary → State MFA → UAE Embassy → UAE MOFA.",
            "india": "India: MEA attestation → UAE Embassy → UAE MOFA.",
            "luxembourg": "Luxembourg: Apostille (Hague) → UAE Embassy → UAE MOFA.",
            "iran": "Iran: No Apostille — full MFA chain → UAE Embassy Tehran → UAE MOFA.",
            "china": "China: CCPIT/Notary → MFA China → UAE Embassy → UAE MOFA.",
            "pakistan": "Pakistan: MOFA Pakistan → UAE Embassy → UAE MOFA.",
            "saudi arabia": "Saudi Arabia: MFA KSA → UAE Embassy Riyadh → UAE MOFA.",
        }
        for entry in corp_kyc:
            ent_name = _s(entry.get("entity"))
            ent_pct  = _s(entry.get("share_pct")) or "—"
            ent_jur  = _s(entry.get("jurisdiction")) or "—"
            _sub_header(doc, f"{ent_name} — {ent_pct} — {ent_jur}")

            provided_set = set(entry.get("provided") or [])
            req_rows: list[tuple[str, str]] = []
            for d in entry.get("required_docs") or []:
                tick = "✓ Provided" if d in provided_set else "✗ Missing"
                req_rows.append((d, tick))
            _kv_table(doc, req_rows)

            att = entry.get("attestation") or {}
            def _stage_sym(v):
                if v is True:  return "✓"
                if v is False: return "✗"
                return "—"
            s2 = att.get("stage2_home_country") or {}
            s2_summary = (
                f"Notary {_stage_sym(s2.get('notary'))} | "
                f"MFA {_stage_sym(s2.get('mfa'))} | "
                f"Apostille {_stage_sym(s2.get('apostille'))}"
            )
            att_rows = [
                ("Stage 1 — Translation",
                 _stage_sym((att.get("stage1_translation") or {}).get("present"))),
                ("Stage 2 — Home Country", s2_summary),
                ("Stage 3 — UAE Embassy",
                 _stage_sym((att.get("stage3_uae_embassy") or {}).get("present"))),
                ("Stage 4 — UAE MOFA",
                 _stage_sym((att.get("stage4_uae_mofa") or {}).get("present"))),
            ]
            _kv_table(doc, att_rows)

            jur_lc = ent_jur.lower()
            note = next((v for k, v in _country_notes.items() if k in jur_lc), None)
            if note:
                _kv_table(doc, [("Country attestation note", note)])

    # ══════════════════════════════════════════════════════════════════════════
    # N — ADDRESS VERIFICATION — CROSS-DOCUMENT
    # ══════════════════════════════════════════════════════════════════════════
    any_addr_docs = (tl_up and ej_up) or (tl_up and vat_up) or (ej_up and vat_up)
    if any_addr_docs:
        _section_header(doc, n, "Address Verification — Cross-Document"); n += 1

        tl_unit  = _v(tl, "unit_number",        "")
        tl_bldg  = _v(tl, "building_name",      "")
        tl_area  = _v(tl, "area",               "")
        tl_parc  = _v(tl, "parcel_id",          "")
        tl_addr  = _v(tl, "registered_address", "")
        ej_unit  = _v(ej, "unit_number",         "")
        ej_bldg  = _v(ej, "building_name",       "")
        ej_area  = _v(ej, "area",                "")
        ej_parc  = _v(ej, "land_dm_parcel_id",   "")
        ej_tenant  = _v(ej, "tenant_name",      "")
        tl_company = _v(tl, "company_name",     "")
        ej_lic_no  = _v(ej, "licence_number",   "")
        tl_lic_no  = _v(tl, "license_number",   "")
        vat_addr   = _v(vat, "registered_address", "")

        if vat_up and (tl_up or ej_up):
            # 3-way table
            addr3: list[tuple] = [
                ("Company / Tenant Name",
                 tl_company or "—", ej_tenant or "—", _v(vat, "company_name") or "—",
                 _match3(tl_company, ej_tenant, _v(vat, "company_name", ""))),
                ("Full Address",
                 tl_addr or "—", (ej_bldg + " " + ej_area).strip() or "—", vat_addr or "—",
                 _match3(tl_addr, (ej_bldg + " " + ej_area).strip(), vat_addr)),
                ("Unit No.",
                 tl_unit or "—", ej_unit or "—", "—",
                 _match2(tl_unit, ej_unit)),
                ("Building",
                 tl_bldg or "—", ej_bldg or "—", "—",
                 _match2(tl_bldg, ej_bldg)),
                ("Area",
                 tl_area or "—", ej_area or "—", "—",
                 _match2(tl_area, ej_area)),
                ("Parcel ID / Land DM",
                 tl_parc or "—", ej_parc or "—", "—",
                 _match2(tl_parc, ej_parc)),
                ("Licence No.",
                 tl_lic_no or "—", ej_lic_no or "—", "—",
                 _match2(tl_lic_no, ej_lic_no)),
            ]
            _three_way_table(doc,
                ("Field", "Trade Licence", "EJARI", "VAT Certificate", "Match"),
                addr3)
            if vat_addr and tl_addr and _match2(vat_addr, tl_addr) != "✓":
                _flag("VAT Address Mismatch",
                      "VAT Certificate vs Trade Licence",
                      "Registered Address", tl_addr, vat_addr,
                      "Client must update address with the Federal Tax Authority (FTA) "
                      "to match current Trade Licence address. UAE VAT compliance requirement.")
        else:
            # 2-way EJARI vs TL
            cross: list[tuple] = [
                ("Company / Tenant Name",
                 ej_tenant  or "—", tl_company or "—",
                 _match2(ej_tenant, tl_company)),
                ("Licence No.",
                 ej_lic_no  or "—", tl_lic_no  or "—",
                 _match2(ej_lic_no, tl_lic_no)),
                ("Unit No.",
                 ej_unit    or "—", tl_unit    or "—",
                 _match2(ej_unit, tl_unit)),
                ("Building",
                 ej_bldg    or "—", tl_bldg    or "—",
                 _match2(ej_bldg, tl_bldg)),
                ("Area",
                 ej_area    or "—", tl_area    or "—",
                 _match2(ej_area, tl_area)),
                ("Parcel ID / Land DM",
                 ej_parc    or "—", tl_parc    or "—",
                 _match2(ej_parc, tl_parc)),
            ]
            _verify_table(doc, ("Field", "EJARI", "Trade Licence", "Match"), cross)
            for f, ev, tv, m in cross:
                if m == "⚠":
                    _flag(f"Mismatch — {f}", "EJARI vs Trade Licence", f, ev, tv,
                          "Verify with client and update the relevant document.")

    # ══════════════════════════════════════════════════════════════════════════
    # 13 — NAME VERIFICATION — TRADE LICENCE vs MOA
    # ══════════════════════════════════════════════════════════════════════════
    nv_rows: list[tuple] = []
    if tl_up and moa_up:
        moa_shareholders = moa.get("shareholders") if isinstance(moa.get("shareholders"), list) else []
        moa_managers     = moa.get("managers")     if isinstance(moa.get("managers"), list)     else []

        def _join_names(items, key):
            return ", ".join(_s(x.get(key, "")).strip() for x in items
                             if isinstance(x, dict) and x.get(key)) or ""

        tl_own  = _v(tl,  "owner_name",          "")
        moa_own = _join_names(moa_shareholders, "name") or _v(moa, "owner_name", "")
        tl_onat = _v(tl,  "owner_nationality",   "")
        moa_onat= (moa_shareholders[0].get("nationality") if moa_shareholders else "") \
                 or _v(moa, "owner_nationality", "")
        moa_onat= _s(moa_onat).strip()
        tl_oshr = _v(tl,  "owner_share",         "")
        moa_oshr= (moa_shareholders[0].get("shares") or moa_shareholders[0].get("share_percentage")
                   if moa_shareholders else "") or _v(moa, "owner_shares", "")
        moa_oshr= _s(moa_oshr).strip()
        tl_mgr  = _v(tl,  "manager_name",        "")
        moa_mgr = _join_names(moa_managers, "name") or _v(moa, "manager_name", "")
        tl_mnat = _v(tl,  "manager_nationality", "")
        moa_mnat= (moa_managers[0].get("nationality") if moa_managers else "") \
                 or _v(moa, "manager_nationality", "")
        moa_mnat= _s(moa_mnat).strip()
        tl_mrol = _v(tl,  "manager_role",        "")
        moa_mrol= (moa_managers[0].get("role") if moa_managers else "") \
                 or _v(moa, "manager_role", "") or _v(moa, "signing_authority", "")
        moa_mrol= _s(moa_mrol).strip()

        # Multi-name match: split TL comma-separated values and check token-set overlap
        def _multi_name_match(a: str, b: str) -> str:
            if not a or a == "—" or not b or b == "—":
                return "—"
            a_parts = [p.strip() for p in str(a).split(",") if p.strip()]
            b_parts = [p.strip() for p in str(b).split(",") if p.strip()]
            if len(a_parts) <= 1 and len(b_parts) <= 1:
                return _match2(a, b)
            # Each name on the smaller side must match some name on the other side
            small, large = (a_parts, b_parts) if len(a_parts) <= len(b_parts) else (b_parts, a_parts)
            unmatched = []
            for sname in small:
                if not any(_names_match(sname, lname) for lname in large):
                    unmatched.append(sname)
            return "✓" if not unmatched else "⚠"

        nv_rows = [
            ("Owner Name",           tl_own  or "—", moa_own  or "—", _multi_name_match(tl_own,  moa_own)),
            ("Owner Nationality",    tl_onat or "—", moa_onat or "—", _match2(tl_onat, moa_onat)),
            ("Owner Share",          tl_oshr or "—", moa_oshr or "—", _percent_match(tl_oshr, moa_oshr)),
            ("Manager Name",         tl_mgr  or "—", moa_mgr  or "—", _multi_name_match(tl_mgr,  moa_mgr)),
            ("Manager Nationality",  tl_mnat or "—", moa_mnat or "—", _match2(tl_mnat, moa_mnat)),
            ("Manager Role",         tl_mrol or "—", moa_mrol or "—", _match2(tl_mrol, moa_mrol)),
        ]
        _section_header(doc, n, "Name Verification — Trade Licence vs MOA"); n += 1
        _verify_table(doc, ("Field", "Trade Licence", "MOA", "Match"), nv_rows)

        _name_field_actions = {
            "Owner Name":          ("Name Mismatch",          "Align owner name spelling across Trade Licence and MOA."),
            "Owner Nationality":   ("Nationality Mismatch",   "Verify owner nationality on Trade Licence and MOA."),
            "Owner Share":         ("Share Mismatch",         "Verify shareholding percentage on Trade Licence and MOA."),
            "Manager Name":        ("Name Mismatch",          "Align manager name spelling across Trade Licence and MOA."),
            "Manager Nationality": ("Nationality Mismatch",   "Verify manager nationality on Trade Licence and MOA."),
            "Manager Role":        ("Role Mismatch",          "Verify manager role on Trade Licence and MOA."),
        }
        for f, tv, mv, m in nv_rows:
            if m == "⚠":
                ftype, action = _name_field_actions.get(f, ("Mismatch", "Verify and align values across Trade Licence and MOA."))
                _flag(f"{ftype} — {f}", "Trade Licence vs MOA", f, tv, mv, action)

    # ══════════════════════════════════════════════════════════════════════════
    # 14 — PERSONAL DOCUMENTS VERIFICATION
    # ══════════════════════════════════════════════════════════════════════════
    _any_personal = eid_up or pp_up or visa_up
    if has_multi_partner:
        _any_personal = any(
            (pdoc.get("passport") or pdoc.get("emirates_id") or pdoc.get("residence_visa"))
            for pdoc in partner_docs_list
        )

    if _any_personal:
        _section_header(doc, n, "Personal Documents Verification"); n += 1

        if has_multi_partner:
            # ── Multi-partner verification ────────────────────────────────
            all_status_rows: list[tuple] = []
            for pidx, pdoc in enumerate(partner_docs_list, 1):
                pname = pdoc.get("partner_name", f"Partner {pidx}")
                ppp   = dict(pdoc.get("passport") or {})
                peid  = dict(pdoc.get("emirates_id") or {})
                pvisa = dict(pdoc.get("residence_visa") or {})
                ppp.pop("error", None); peid.pop("error", None); pvisa.pop("error", None)
                if not (ppp or peid or pvisa):
                    continue

                # Doc validity rows (combined table)
                if peid:
                    peid_lbl, peid_sym = _expiry_label(peid.get("expiry_date"), today)
                    all_status_rows.append((
                        f"Emirates ID — {pname}", _v(peid, "holder_name"),
                        _v(peid, "id_number"), _fmt_date(peid.get("expiry_date")), peid_sym))
                    if peid_sym in ("⚠", "✗"):
                        _flag(f"Emirates ID Validity — {pname}", "Emirates ID",
                              "Expiry Date", peid_lbl, "",
                              "Renew Emirates ID immediately." if peid_sym == "✗"
                              else "Emirates ID expiring within 30 days.")
                if ppp:
                    ppp_lbl, ppp_sym = _expiry_label(ppp.get("expiry_date"), today)
                    all_status_rows.append((
                        f"Passport — {pname}", _v(ppp, "holder_name"),
                        _v(ppp, "passport_number"), _fmt_date(ppp.get("expiry_date")), ppp_sym))
                    if ppp_sym in ("⚠", "✗"):
                        _flag(f"Passport Validity — {pname}", "Passport",
                              "Expiry Date", ppp_lbl, "",
                              "Renew passport immediately." if ppp_sym == "✗"
                              else "Passport expiring within 30 days.")
                if pvisa:
                    pvisa_lbl, pvisa_sym = _expiry_label(pvisa.get("expiry_date"), today)
                    all_status_rows.append((
                        f"Residence Visa — {pname}", _v(pvisa, "holder_name"),
                        _v(pvisa, "visa_number"), _fmt_date(pvisa.get("expiry_date")), pvisa_sym))
                    if pvisa_sym in ("⚠", "✗"):
                        _flag(f"Residence Visa Validity — {pname}", "UAE Residence Visa",
                              "Expiry Date", pvisa_lbl, "",
                              "Renew residence visa immediately." if pvisa_sym == "✗"
                              else "Visa expiring within 30 days.")

            if all_status_rows:
                _sub_header(doc, "14A — Document Validity Status")
                _doc_status_table(doc, all_status_rows)

            # Per-partner cross-verification
            co_name = _v(tl, "company_name", "") or _v(moa, "company_name", "")
            for pidx, pdoc in enumerate(partner_docs_list, 1):
                pname = pdoc.get("partner_name", f"Partner {pidx}")
                ppp   = dict(pdoc.get("passport") or {})
                peid  = dict(pdoc.get("emirates_id") or {})
                pvisa = dict(pdoc.get("residence_visa") or {})
                ppp.pop("error", None); peid.pop("error", None); pvisa.pop("error", None)
                p_eid_name  = _v(peid,  "holder_name", "") if peid  else ""
                p_pp_name   = _v(ppp,   "holder_name", "") if ppp   else ""
                p_visa_name = _v(pvisa, "holder_name", "") if pvisa else ""
                p_present = [x for x in [p_eid_name, p_pp_name, p_visa_name] if x and x != "—"]

                if len(p_present) >= 2:
                    m3 = _match3(p_eid_name, p_pp_name, p_visa_name)
                    _sub_header(doc, f"14B — Name Cross-Match — {pname}")
                    _three_way_table(doc,
                        ("Field", "Emirates ID", "Passport", "UAE Residence Visa", "Match"),
                        [("Full Name", p_eid_name or "—", p_pp_name or "—",
                          p_visa_name or "—", m3)])
                    if m3 == "⚠":
                        _flag(f"Name Mismatch — {pname}",
                              "EID / Passport / Visa", "Holder Name",
                              f"EID: {p_eid_name} | PP: {p_pp_name} | Visa: {p_visa_name}", "",
                              "Names must be consistent across all personal identity documents.")

                if ppp.get("passport_number") and pvisa.get("passport_number"):
                    pp_no   = _v(ppp,   "passport_number")
                    visa_pp = _v(pvisa, "passport_number")
                    mpp = _match2(pp_no, visa_pp)
                    _sub_header(doc, f"14C — Passport Number — {pname}")
                    _verify_table(doc,
                        ("Field", "Passport", "UAE Residence Visa", "Match"),
                        [("Passport No.", pp_no, visa_pp, mpp)])
                    if mpp == "⚠":
                        _flag(f"Passport No. Mismatch — {pname}",
                              "Passport vs Residence Visa", "Passport No.",
                              pp_no, visa_pp, "Verify passport number.")

                # Name vs partner name from Partners Annex
                if p_present and pname and pname != "—":
                    corp_rows: list[tuple] = []
                    for lbl, nm in [("Emirates ID", p_eid_name), ("Passport", p_pp_name),
                                    ("Residence Visa", p_visa_name)]:
                        if nm and nm != "—":
                            m = _match2(nm, pname)
                            corp_rows.append((f"{lbl} vs Partners Annex", nm, pname, m))
                            if m == "⚠":
                                _flag(f"Name Mismatch — {lbl} vs Partners Annex ({pname})",
                                      f"{lbl} vs Partners Annex", "Holder Name",
                                      nm, pname, "Verify name spelling.")
                    if corp_rows:
                        _sub_header(doc, f"14D — Name vs Partners Annex — {pname}")
                        _verify_table(doc,
                            ("Comparison", "Personal Document", "Partners Annex Name", "Match"),
                            corp_rows)

                if pvisa.get("employer") and co_name and co_name != "—":
                    emp = _v(pvisa, "employer", "")
                    m = _match2(emp, co_name)
                    _sub_header(doc, f"14E — Employer vs Company — {pname}")
                    _verify_table(doc,
                        ("Field", "Visa — Employer / Sponsor", "Company Name (TL / MOA)", "Match"),
                        [("Employer / Company", emp, co_name, m)])
                    if m == "⚠":
                        _flag(f"Employer Mismatch — {pname}",
                              "Residence Visa vs TL/MOA", "Employer / Sponsor",
                              emp, co_name, "Verify sponsor on visa matches company name.")

        else:
            # ── Single-partner verification (original logic) ──────────────
            _sub_header(doc, "14A — Document Validity Status")
            status_rows: list[tuple] = []
            if eid_up and eid:
                status_rows.append((
                    "Emirates ID", _v(eid, "holder_name"), _v(eid, "id_number"),
                    _fmt_date(eid.get("expiry_date")), eid_sym))
                if eid_sym in ("⚠", "✗"):
                    _flag("Emirates ID Validity", "Emirates ID", "Expiry Date", eid_label, "",
                          "Renew Emirates ID immediately." if eid_sym == "✗"
                          else "Emirates ID expiring within 30 days.")
            if pp_up and pp:
                status_rows.append((
                    "Passport", _v(pp, "holder_name"), _v(pp, "passport_number"),
                    _fmt_date(pp.get("expiry_date")), pp_sym))
                if pp_sym in ("⚠", "✗"):
                    _flag("Passport Validity", "Passport", "Expiry Date", pp_label, "",
                          "Renew passport immediately." if pp_sym == "✗"
                          else "Passport expiring within 30 days.")
            if visa_up and visa:
                status_rows.append((
                    "UAE Residence Visa", _v(visa, "holder_name"), _v(visa, "visa_number"),
                    _fmt_date(visa.get("expiry_date")), visa_sym))
                if visa_sym in ("⚠", "✗"):
                    _flag("Residence Visa Validity", "UAE Residence Visa", "Expiry Date",
                          visa_label, "",
                          "Renew residence visa immediately." if visa_sym == "✗"
                          else "Visa expiring within 30 days.")
            _doc_status_table(doc, status_rows)

            eid_name  = (_v(eid,  "holder_name", "") if eid_up  and eid  else "")
            pp_name   = (_v(pp,   "holder_name", "") if pp_up   and pp   else "")
            visa_name = (_v(visa, "holder_name", "") if visa_up and visa else "")
            present_names = [x for x in [eid_name, pp_name, visa_name] if x and x != "—"]
            if len(present_names) >= 2:
                _sub_header(doc, "14B — Name Cross-Match — Personal Documents")
                m3 = _match3(eid_name, pp_name, visa_name)
                _three_way_table(doc,
                    ("Field", "Emirates ID", "Passport", "UAE Residence Visa", "Match"),
                    [("Full Name", eid_name or "—", pp_name or "—", visa_name or "—", m3)])
                if m3 == "⚠":
                    _flag("Personal Document Name Mismatch",
                          "Emirates ID vs Passport vs UAE Residence Visa", "Holder Name",
                          f"EID: {eid_name} | Passport: {pp_name} | Visa: {visa_name}", "",
                          "Names must be consistent across all personal identity documents.")

            if pp_up and visa_up and pp.get("passport_number") and visa.get("passport_number"):
                _sub_header(doc, "14C — Passport Number Consistency")
                pp_no   = _v(pp,   "passport_number")
                visa_pp = _v(visa, "passport_number")
                mpp = _match2(pp_no, visa_pp)
                _verify_table(doc,
                    ("Field", "Passport", "UAE Residence Visa", "Match"),
                    [("Passport No.", pp_no, visa_pp, mpp)])
                if mpp == "⚠":
                    _flag("Passport Number Mismatch", "Passport vs UAE Residence Visa",
                          "Passport No.", pp_no, visa_pp,
                          "Verify passport number — visa may reference an expired passport.")

            ref_name = (
                _v(moa, "owner_name",   "") or _v(tl, "owner_name",   "") or
                _v(moa, "manager_name", "") or _v(tl, "manager_name", "")
            )
            if ref_name and ref_name != "—" and present_names:
                corp_rows: list[tuple] = []
                if eid_name and eid_name != "—":
                    m = _match2(eid_name, ref_name)
                    corp_rows.append(("Emirates ID vs TL / MOA", eid_name, ref_name, m))
                    if m == "⚠":
                        _flag("Name Mismatch — EID vs Corporate",
                              "Emirates ID vs Trade Licence / MOA", "Holder Name",
                              eid_name, ref_name, "Verify name spelling.")
                if pp_name and pp_name != "—":
                    m = _match2(pp_name, ref_name)
                    corp_rows.append(("Passport vs TL / MOA", pp_name, ref_name, m))
                    if m == "⚠":
                        _flag("Name Mismatch — Passport vs Corporate",
                              "Passport vs Trade Licence / MOA", "Holder Name",
                              pp_name, ref_name, "Verify name spelling.")
                if visa_name and visa_name != "—":
                    m = _match2(visa_name, ref_name)
                    corp_rows.append(("Residence Visa vs TL / MOA", visa_name, ref_name, m))
                    if m == "⚠":
                        _flag("Name Mismatch — Visa vs Corporate",
                              "Residence Visa vs Trade Licence / MOA", "Holder Name",
                              visa_name, ref_name, "Verify name spelling.")
                if corp_rows:
                    _sub_header(doc, "14D — Name vs Corporate Documents")
                    _verify_table(doc,
                        ("Comparison", "Personal Document", "TL / MOA Name", "Match"),
                        corp_rows)

            if visa_up and visa.get("employer"):
                co  = _v(tl, "company_name", "") or _v(moa, "company_name", "")
                emp = _v(visa, "employer", "")
                if co and co != "—":
                    _sub_header(doc, "14E — Employer (Visa) vs Company Name")
                    m = _match2(emp, co)
                    _verify_table(doc,
                        ("Field", "Visa — Employer / Sponsor", "Company Name (TL / MOA)", "Match"),
                        [("Employer / Company", emp, co, m)])
                    if m == "⚠":
                        _flag("Employer Mismatch", "UAE Residence Visa vs Trade Licence / MOA",
                              "Employer / Sponsor", emp, co,
                              "Verify sponsor on visa matches the company name on Trade Licence.")

    # ══════════════════════════════════════════════════════════════════════════
    # 18 — KYC VERIFICATION CHECKLIST (analysis-driven, A–G)
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, n, "KYC Verification Checklist"); n += 1

    _CHECKLIST_GROUPS = [
        ("A", "Corporate Documents"),
        ("B", "Personal Documents"),
        ("C", "Cross-Verification"),
        ("D", "Banking & Authority"),
        ("E", "Physical Presence & POA"),
        ("F", "Corporate Shareholder KYC"),
        ("G", "Final Status"),
    ]
    _STATUS_SYM = {"pass": "☑", "warn": "⚠", "fail": "☐", "na": "—"}

    checklist_dict = (analysis.get("checklist") or {}) if isinstance(analysis, dict) else {}
    if checklist_dict:
        for letter, group_label in _CHECKLIST_GROUPS:
            items = checklist_dict.get(letter) or []
            if not items:
                continue
            _sub_header(doc, f"{letter} — {group_label}")
            rows: list[tuple[str, str]] = []
            for it in items:
                sym = _STATUS_SYM.get(it.get("status", ""), "—")
                detail = _s(it.get("detail")) or ""
                rows.append((f"{sym}  {_s(it.get('label'))}", detail))
            _kv_table(doc, rows)
    else:
        _kv_table(doc, [("Checklist", "Not provided.")])

    # ══════════════════════════════════════════════════════════════════════════
    # 19 — DISCREPANCIES & FLAGS
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, n, "Discrepancies & Flags"); n += 1
    analysis_flags = (analysis.get("flags") or []) if isinstance(analysis, dict) else []
    if not analysis_flags:
        _kv_table(doc, [("Result",
                          "✓  No discrepancies identified. "
                          "All documents reviewed are consistent.")])
    else:
        for i, flag in enumerate(analysis_flags, 1):
            sev = (flag.get("severity") or "").lower()
            icon = "❌" if sev == "error" else ("⚠️" if sev == "warn" else "•")
            code = _s(flag.get("code")) or "FLAG"
            _sub_header(doc, f"{icon}  FLAG {i}: {code}")
            docs_affected = flag.get("documents_affected") or []
            if isinstance(docs_affected, (list, tuple)):
                docs_str = ", ".join(_s(x) for x in docs_affected if _s(x))
            else:
                docs_str = _s(docs_affected)
            flag_rows: list[tuple] = []
            if docs_str:
                flag_rows.append(("Documents Affected", docs_str))
            if flag.get("field"):
                flag_rows.append(("Field",              _s(flag.get("field"))))
            if flag.get("issue"):
                flag_rows.append(("Issue",              _s(flag.get("issue"))))
            if flag.get("recommended_action"):
                flag_rows.append(("Recommended Action", _s(flag.get("recommended_action"))))
            if flag.get("kyc_status"):
                flag_rows.append(("KYC Status",         _s(flag.get("kyc_status"))))
            _kv_table(doc, flag_rows)

    # ══════════════════════════════════════════════════════════════════════════
    # 20 — DOCUMENTS REVIEWED
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, n, "Documents Reviewed"); n += 1
    reviewed: list[tuple] = []
    idx = 1
    if moa:
        line = f"Memorandum of Association — Contract No. {_v(moa, 'contract_number', '—')}"
        if moa.get("moa_date"):
            line += f", dated {_fmt_date(moa['moa_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if tl:
        line = f"Trade Licence No. {license_no} — {_v(tl, 'issuing_authority', 'DET Dubai')}"
        if tl.get("expiry_date"):
            line += f", expiry {_fmt_date(tl['expiry_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if ej:
        line = f"EJARI Tenancy Contract No. {ejari_no}"
        if ej.get("registration_date"):
            line += f" — registered {_fmt_date(ej['registration_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if vat:
        line = f"VAT Registration Certificate — TRN {_v(vat, 'trn', '—')}"
        if vat.get("effective_date"):
            line += f", effective {_fmt_date(vat['effective_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if ins:
        line = f"Insurance Certificate — {_v(ins, 'insurer', '—')}"
        if ins.get("policy_number"):
            line += f", Policy No. {_v(ins, 'policy_number')}"
        reviewed.append((str(idx), line)); idx += 1
    if tl and tl.get("register_number"):
        reviewed.append((str(idx),
            f"Commercial Register Certificate — No. {_v(tl, 'register_number')}")); idx += 1
    if tl and tl.get("last_renewal_fee"):
        line = f"Renewal Receipt — {_v(tl, 'last_renewal_fee')}"
        if tl.get("last_renewal_date"):
            line += f" paid {_fmt_date(tl['last_renewal_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if pp:
        line = f"Passport — {_v(pp, 'holder_name')}, No. {_v(pp, 'passport_number')}"
        if pp.get("expiry_date"):
            line += f", expiry {_fmt_date(pp['expiry_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if eid:
        line = f"Emirates ID — {_v(eid, 'holder_name')}, No. {_v(eid, 'id_number')}"
        if eid.get("expiry_date"):
            line += f", expiry {_fmt_date(eid['expiry_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if visa:
        line = f"UAE Residence Visa — {_v(visa, 'holder_name')}"
        if visa.get("visa_number"):
            line += f", Permit No. {_v(visa, 'visa_number')}"
        if visa.get("expiry_date"):
            line += f", expiry {_fmt_date(visa['expiry_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if br:
        line = f"Board Resolution — {_v(br, 'resolution_type', 'Resolution')}"
        if br.get("resolution_date"):
            line += f", dated {_fmt_date(br['resolution_date'])}"
        if br.get("signatory_name"):
            line += f", signatory: {_v(br, 'signatory_name')}"
        reviewed.append((str(idx), line)); idx += 1
    if poa:
        line = f"Power of Attorney — Grantor: {_v(poa, 'grantor_name', '—')}"
        line += f", Grantee: {_v(poa, 'grantee_name', '—')}"
        if poa.get("poa_date"):
            line += f", dated {_fmt_date(poa['poa_date'])}"
        reviewed.append((str(idx), line)); idx += 1
    if pa:
        line = "Partners Annex / Schedule of Partners"
        if pa.get("company_name"):
            line += f" — {_v(pa, 'company_name')}"
        partners_list = pa.get("partners", [])
        if isinstance(partners_list, list):
            line += f", {len(partners_list)} partner(s) listed"
        reviewed.append((str(idx), line)); idx += 1
    _kv_table(doc, reviewed)

    # ── DISCLAIMER ─────────────────────────────────────────────────────────────
    _spacer(doc, 14)
    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(4)
    dr = dp.add_run("DISCLAIMER")
    dr.bold = True
    dr.font.size = Pt(9)
    dr.font.color.rgb = _rgb(_NAVY)

    bp = doc.add_paragraph()
    bp.paragraph_format.space_after = Pt(0)
    disc_run = bp.add_run(
        "This KYC Profile has been prepared by National Assurance & Advisory Services "
        "FZ LLC (NAAS), Office 319, Garhoud Star Building, Al Garhoud, Dubai, UAE, "
        "from documents supplied by the client. This document does not constitute "
        "legal advice, a credit opinion, or a regulatory compliance clearance. NAAS "
        "accepts no liability for actions taken in reliance hereon without independent "
        "verification of the original source documents."
    )
    disc_run.italic = True
    disc_run.font.size = Pt(8)
    disc_run.font.color.rgb = _rgb(_GREY_LITE)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Report data builder for frontend display ──────────────────────────────────

def build_report_data(extracted: dict, today: date, analysis: dict | None = None) -> dict:
    """
    Build structured report data for the frontend preview page.
    Returns { company_name, panels, flags, analysis? }.

    Panel types:
      "docstatus" — document validity table
      "kv"        — key/value detail table
      "match"     — cross-document comparison table

    `analysis` is the optional NAAS v4.0 compliance dict (validity / cross_checks /
    moa_authority / presence / shareholders / corporate_kyc / checklist / flags /
    version) produced by `app.kyc_compliance.analyse`. It is stashed verbatim so
    the frontend can read it without re-parsing the DOCX.
    """

    # ── helpers ────────────────────────────────────────────────────────────────
    def _msym(s: str) -> str:
        """Map ✓/⚠/— to ok/warn/dash for the frontend."""
        if s == "✓": return "ok"
        if s == "⚠": return "warn"
        return "dash"

    def _sc(s: str) -> str:
        if s == "✓": return "valid"
        if s == "⚠": return "expiring_soon"
        if s == "✗": return "expired"
        return "no_expiry"

    def _st(s: str, active: str = "VALID") -> str:
        if s == "✓": return active
        if s == "⚠": return "EXPIRING SOON"
        if s == "✗": return "EXPIRED"
        return "N/A"

    def _kv(label: str, value: str, sym: str = "") -> dict:
        return {"label": label, "value": value, "sym": sym}

    def _mr(field: str, values: list, match: str) -> dict:
        return {"field": field, "values": values, "match": match}

    # ── unpack ─────────────────────────────────────────────────────────────────
    tl   = dict(extracted.get("trade_license")   or {})
    ej   = dict(extracted.get("ejari")            or {})
    moa  = dict(extracted.get("moa")              or {})
    ins  = dict(extracted.get("insurance")        or {})
    pp   = dict(extracted.get("passport")         or {})
    eid  = dict(extracted.get("emirates_id")      or {})
    visa = dict(extracted.get("residence_visa")   or {})
    vat  = dict(extracted.get("vat_certificate")  or {})
    br   = dict(extracted.get("board_resolution") or {})
    poa  = dict(extracted.get("poa")              or {})
    pa   = dict(extracted.get("partners_annex")   or {})
    for d in [tl, ej, moa, ins, pp, eid, visa, vat, br, poa, pa]:
        d.pop("error", None)

    tl_up   = bool(tl);  ej_up  = bool(ej);  moa_up  = bool(moa); ins_up  = bool(ins)
    pp_up   = bool(pp);  eid_up = bool(eid); visa_up = bool(visa); vat_up  = bool(vat)
    br_up   = bool(br);  poa_up = bool(poa); pa_up   = bool(pa)

    tl_lbl,   tl_sym   = _expiry_label(tl.get("expiry_date"),   today)
    ej_lbl,   ej_sym   = _expiry_label(ej.get("expiry_date"),   today)
    pp_lbl,   pp_sym   = _expiry_label(pp.get("expiry_date"),   today)
    eid_lbl,  eid_sym  = _expiry_label(eid.get("expiry_date"),  today)
    visa_lbl, visa_sym = _expiry_label(visa.get("expiry_date"), today)
    ins_lbl,  ins_sym  = _insurance_label(ins.get("valid_to"),  today)

    company_name = (_v(tl, "company_name") or _v(moa, "company_name") or _v(vat, "company_name") or "—")

    partner_docs_list = extracted.get("partner_personal_docs") or []
    has_multi_partner = bool(partner_docs_list)

    panels: list[dict] = []
    flags:  list[dict] = []

    def _flag(ftype, docs, field, val_a, val_b, action):
        flags.append({"type": ftype, "docs": docs, "field": field,
                      "val_a": val_a, "val_b": val_b, "action": action})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 1 — Document Validity Status
    # ══════════════════════════════════════════════════════════════════════════
    ds_rows: list[dict] = []
    if tl_up:
        ds_rows.append({"doc": "Trade Licence",   "name": _v(tl,  "company_name"),
                        "number": _v(tl,  "license_number"),
                        "expiry": _fmt_date(tl.get("expiry_date")),
                        "sc": _sc(tl_sym), "st": _st(tl_sym), "sym": tl_sym})
        if tl_sym in ("⚠","✗"):
            _flag("Trade Licence Validity", "Trade Licence", "Expiry Date", tl_lbl, "",
                  "Renew Trade Licence immediately." if tl_sym == "✗"
                  else "Initiate renewal — expires within 30 days.")
    if ej_up:
        ds_rows.append({"doc": "EJARI Contract",  "name": _v(ej,  "tenant_name"),
                        "number": _v(ej,  "ejari_number"),
                        "expiry": _fmt_date(ej.get("expiry_date")),
                        "sc": _sc(ej_sym), "st": _st(ej_sym), "sym": ej_sym})
        if ej_sym in ("⚠","✗"):
            _flag("EJARI Validity", "EJARI", "Lease End Date", ej_lbl, "",
                  "Renew EJARI / tenancy contract." if ej_sym == "✗"
                  else "EJARI renewal due within 30 days.")
    if ins_up:
        ds_rows.append({"doc": "Insurance",       "name": _v(ins, "insured_name"),
                        "number": _v(ins, "policy_number"),
                        "expiry": _fmt_date(ins.get("valid_to")),
                        "sc": _sc(ins_sym), "st": _st(ins_sym, "ACTIVE"), "sym": ins_sym})
        if ins_sym in ("⚠","✗"):
            _flag("Insurance Validity", "Insurance Certificate", "Valid To", ins_lbl, "",
                  "Renew insurance policy." if ins_sym == "✗"
                  else "Insurance expiring within 30 days.")
    if has_multi_partner:
        # Add per-partner personal doc rows to the status table
        for pidx, pdoc in enumerate(partner_docs_list, 1):
            pname = pdoc.get("partner_name", f"Partner {pidx}")
            for dtype, dkey, nkey in [("Passport", "passport", "passport_number"),
                                       ("Emirates ID", "emirates_id", "id_number"),
                                       ("Residence Visa", "residence_visa", "visa_number")]:
                d = dict(pdoc.get(dkey) or {})
                d.pop("error", None)
                if not d:
                    continue
                d_lbl, d_sym = _expiry_label(d.get("expiry_date"), today)
                ds_rows.append({"doc": f"{dtype} — {pname}", "name": _v(d, "holder_name"),
                                "number": _v(d, nkey),
                                "expiry": _fmt_date(d.get("expiry_date")),
                                "sc": _sc(d_sym), "st": _st(d_sym), "sym": d_sym})
                if d_sym in ("⚠", "✗"):
                    _flag(f"{dtype} Validity — {pname}", dtype, "Expiry Date", d_lbl, "",
                          f"Renew {dtype.lower()} immediately." if d_sym == "✗"
                          else f"{dtype} expiring within 30 days.")
    else:
        if pp_up:
            ds_rows.append({"doc": "Passport",        "name": _v(pp,  "holder_name"),
                            "number": _v(pp,  "passport_number"),
                            "expiry": _fmt_date(pp.get("expiry_date")),
                            "sc": _sc(pp_sym), "st": _st(pp_sym), "sym": pp_sym})
            if pp_sym in ("⚠","✗"):
                _flag("Passport Validity", "Passport", "Expiry Date", pp_lbl, "",
                      "Renew passport immediately." if pp_sym == "✗"
                      else "Passport expiring within 30 days.")
        if eid_up:
            ds_rows.append({"doc": "Emirates ID",     "name": _v(eid, "holder_name"),
                            "number": _v(eid, "id_number"),
                            "expiry": _fmt_date(eid.get("expiry_date")),
                            "sc": _sc(eid_sym), "st": _st(eid_sym), "sym": eid_sym})
            if eid_sym in ("⚠","✗"):
                _flag("Emirates ID Validity", "Emirates ID", "Expiry Date", eid_lbl, "",
                      "Renew Emirates ID immediately." if eid_sym == "✗"
                      else "Emirates ID expiring within 30 days.")
        if visa_up:
            ds_rows.append({"doc": "Residence Visa",  "name": _v(visa,"holder_name"),
                            "number": _v(visa,"visa_number"),
                            "expiry": _fmt_date(visa.get("expiry_date")),
                            "sc": _sc(visa_sym), "st": _st(visa_sym), "sym": visa_sym})
            if visa_sym in ("⚠","✗"):
                _flag("Residence Visa Validity", "UAE Residence Visa", "Expiry Date", visa_lbl, "",
                      "Renew residence visa immediately." if visa_sym == "✗"
                      else "Visa expiring within 30 days.")
    if moa_up:
        ds_rows.append({"doc": "MOA",             "name": _v(moa, "company_name"),
                        "number": _v(moa, "contract_number"),
                        "expiry": "—", "sc": "no_expiry", "st": "N/A", "sym": "—"})
    if vat_up:
        ds_rows.append({"doc": "VAT Certificate", "name": _v(vat, "company_name"),
                        "number": _v(vat, "trn"),
                        "expiry": "—", "sc": "no_expiry", "st": "N/A", "sym": "—"})
    panels.append({"title": "Document Validity Status", "type": "docstatus", "rows": ds_rows})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 2 — Company & Licence Details
    # ══════════════════════════════════════════════════════════════════════════
    if tl_up or moa_up or vat_up:
        kv: list[dict] = []
        cn = _v(tl,"company_name","") or _v(moa,"company_name","") or _v(vat,"company_name","")
        if cn and cn != "—":
            kv.append(_kv("Company Name (English)", cn))
        cn_ar = _v(tl,"company_name_arabic","") or _v(moa,"company_name_arabic","")
        if cn_ar and cn_ar != "—": kv.append(_kv("Company Name (Arabic)", cn_ar))
        lf = _v(tl,"legal_form","") or _v(moa,"legal_form","")
        if lf and lf != "—":   kv.append(_kv("Legal Type", lf))
        if tl.get("license_number"):     kv.append(_kv("Trade Licence No.", _v(tl,"license_number")))
        if tl.get("register_number"):    kv.append(_kv("Commercial Register No.", _v(tl,"register_number")))
        if tl.get("dcci_membership_number"): kv.append(_kv("DCCI Membership No.", _v(tl,"dcci_membership_number")))
        if tl.get("issuing_authority"):  kv.append(_kv("Issuing Authority", _v(tl,"issuing_authority")))
        if tl.get("license_type"):       kv.append(_kv("Licence Type", _v(tl,"license_type")))
        if tl.get("issue_date"):         kv.append(_kv("Issue Date", _fmt_date(tl["issue_date"])))
        if tl.get("expiry_date"):        kv.append(_kv("Licence Expiry", tl_lbl, tl_sym))
        if tl.get("last_renewal_date"):  kv.append(_kv("Last Renewal Date", _fmt_date(tl["last_renewal_date"])))
        if tl.get("last_renewal_fee"):   kv.append(_kv("Last Renewal Fee", _v(tl,"last_renewal_fee")))
        if tl.get("business_activity"):  kv.append(_kv("Business Activity", _v(tl,"business_activity")))
        if tl.get("activity_scope"):     kv.append(_kv("Activity Scope", _v(tl,"activity_scope")))
        if tl.get("registered_address"): kv.append(_kv("Registered Address", _v(tl,"registered_address")))
        if tl.get("unit_number"):        kv.append(_kv("Unit No.", _v(tl,"unit_number")))
        if tl.get("building_name"):      kv.append(_kv("Building", _v(tl,"building_name")))
        if tl.get("area"):               kv.append(_kv("Area", _v(tl,"area")))
        if tl.get("parcel_id"):          kv.append(_kv("Parcel ID / Land DM No.", _v(tl,"parcel_id")))
        if tl.get("makani_number"):      kv.append(_kv("Makani No.", _v(tl,"makani_number")))
        if tl.get("phone_fax"):          kv.append(_kv("Phone / Fax", _v(tl,"phone_fax")))
        if tl.get("mobile"):             kv.append(_kv("Mobile", _v(tl,"mobile")))
        if tl.get("email"):              kv.append(_kv("Email", _v(tl,"email")))
        if vat.get("trn"):               kv.append(_kv("VAT TRN", _v(vat,"trn")))
        if vat.get("effective_date"):    kv.append(_kv("VAT Effective Date", _fmt_date(vat["effective_date"])))
        if moa.get("share_capital"):     kv.append(_kv("Share Capital", _v(moa,"share_capital")))
        if moa.get("shares_count"):      kv.append(_kv("Number of Shares", _v(moa,"shares_count")))
        if moa.get("moa_date"):          kv.append(_kv("MOA Date", _fmt_date(moa["moa_date"])))
        if moa.get("contract_number"):   kv.append(_kv("MOA Contract No.", _v(moa,"contract_number")))
        if moa.get("financial_year"):    kv.append(_kv("Financial Year", _v(moa,"financial_year")))
        if kv: panels.append({"title": "Company & Licence Details", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 3 — Owner & Management
    # ══════════════════════════════════════════════════════════════════════════
    if tl_up or moa_up:
        kv = []
        moa_shareholders_p = moa.get("shareholders") if isinstance(moa.get("shareholders"), list) else []
        moa_shareholders_p = [s for s in moa_shareholders_p if isinstance(s, dict) and s.get("name")]
        moa_managers_p = moa.get("managers") if isinstance(moa.get("managers"), list) else []
        moa_managers_p = [m for m in moa_managers_p if isinstance(m, dict) and m.get("name")]

        if moa_shareholders_p:
            for idx, sh in enumerate(moa_shareholders_p, 1):
                tag = f" {idx}" if len(moa_shareholders_p) > 1 else ""
                kv.append(_kv(f"Shareholder{tag} Name", _s(sh.get("name", "")).strip() or "—"))
                if sh.get("name_arabic"):
                    kv.append(_kv(f"Shareholder{tag} Name (Arabic)", _s(sh["name_arabic"]).strip()))
                if sh.get("nationality"):
                    kv.append(_kv(f"Shareholder{tag} Nationality", _s(sh["nationality"]).strip()))
                shares_val = _s(sh.get("shares", "")).strip() or _s(sh.get("share_percentage", "")).strip()
                if shares_val:
                    kv.append(_kv(f"Shareholder{tag} Shareholding", shares_val))
                if sh.get("person_number"):
                    kv.append(_kv(f"Shareholder{tag} Person No.", _s(sh["person_number"]).strip()))
        else:
            own = _v(moa,"owner_name","") or _v(tl,"owner_name","")
            if own and own != "—":  kv.append(_kv("Owner Name (English)", own))
            own_ar = _v(moa,"owner_name_arabic","")
            if own_ar and own_ar != "—": kv.append(_kv("Owner Name (Arabic)", own_ar))
            own_nat = _v(moa,"owner_nationality","") or _v(tl,"owner_nationality","")
            if own_nat and own_nat != "—": kv.append(_kv("Owner Nationality", own_nat))
            if moa.get("owner_shares"):   kv.append(_kv("Shareholding", _v(moa,"owner_shares")))
            if moa.get("owner_liability"):kv.append(_kv("Liability", _v(moa,"owner_liability")))
            pno = _v(moa,"owner_person_number","") or _v(tl,"owner_person_number","")
            if pno and pno != "—": kv.append(_kv("Person No. (Licence)", pno))

        if moa_managers_p:
            for idx, mg in enumerate(moa_managers_p, 1):
                tag = f" {idx}" if len(moa_managers_p) > 1 else ""
                kv.append(_kv(f"Manager{tag} Name", _s(mg.get("name", "")).strip() or "—"))
                if mg.get("name_arabic"):
                    kv.append(_kv(f"Manager{tag} Name (Arabic)", _s(mg["name_arabic"]).strip()))
                if mg.get("nationality"):
                    kv.append(_kv(f"Manager{tag} Nationality", _s(mg["nationality"]).strip()))
                if mg.get("role"):
                    kv.append(_kv(f"Manager{tag} Role", _s(mg["role"]).strip()))
        else:
            mgr = _v(moa,"manager_name","") or _v(tl,"manager_name","")
            if mgr and mgr != "—":  kv.append(_kv("Manager Name (English)", mgr))
            mgr_ar = _v(moa,"manager_name_arabic","")
            if mgr_ar and mgr_ar != "—": kv.append(_kv("Manager Name (Arabic)", mgr_ar))
            mgr_nat = _v(moa,"manager_nationality","") or _v(tl,"manager_nationality","")
            if mgr_nat and mgr_nat != "—": kv.append(_kv("Manager Nationality", mgr_nat))
            if moa.get("manager_role"):   kv.append(_kv("Manager Role", _v(moa,"manager_role")))

        if moa.get("signing_authority"): kv.append(_kv("Signing Authority", _v(moa,"signing_authority")))
        if moa.get("signing_mode"):   kv.append(_kv("Signing Mode", _v(moa,"signing_mode")))
        if moa.get("authorised_signatory"): kv.append(_kv("Authorised Signatory", _v(moa,"authorised_signatory")))
        if kv: panels.append({"title": "Owner & Management", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 4 — EJARI Details
    # ══════════════════════════════════════════════════════════════════════════
    if ej_up:
        kv = [_kv("EJARI Contract No.", _v(ej,"ejari_number"))]
        if ej.get("registration_date"): kv.append(_kv("Registration Date", _fmt_date(ej["registration_date"])))
        if ej.get("registered_by"):     kv.append(_kv("Registered by", _v(ej,"registered_by")))
        if ej.get("tenant_name"):       kv.append(_kv("Tenant Name", _v(ej,"tenant_name")))
        if ej.get("licence_number"):    kv.append(_kv("Trade Licence No. (EJARI)", _v(ej,"licence_number")))
        if ej.get("licence_issuer"):    kv.append(_kv("Licence Issuer (EJARI)", _v(ej,"licence_issuer")))
        if ej.get("start_date"):        kv.append(_kv("Lease Start Date", _fmt_date(ej["start_date"])))
        if ej.get("expiry_date"):       kv.append(_kv("Lease End Date", ej_lbl, ej_sym))
        if ej.get("annual_rent"):       kv.append(_kv("Annual Rent", _v(ej,"annual_rent")))
        if ej.get("security_deposit"):  kv.append(_kv("Security Deposit", _v(ej,"security_deposit")))
        if ej.get("ejari_fees_paid"):   kv.append(_kv("EJARI Fees Paid", _v(ej,"ejari_fees_paid")))
        if ej.get("unit_number"):       kv.append(_kv("Unit No.", _v(ej,"unit_number")))
        if ej.get("building_name"):     kv.append(_kv("Building", _v(ej,"building_name")))
        if ej.get("area"):              kv.append(_kv("Area", _v(ej,"area")))
        if ej.get("unit_type"):         kv.append(_kv("Unit Type", _v(ej,"unit_type")))
        if ej.get("size"):              kv.append(_kv("Size", _v(ej,"size")))
        if ej.get("plot_number"):       kv.append(_kv("Plot No.", _v(ej,"plot_number")))
        if ej.get("land_dm_parcel_id"): kv.append(_kv("Land DM No. (Parcel ID)", _v(ej,"land_dm_parcel_id")))
        if ej.get("makani_number"):     kv.append(_kv("Makani No.", _v(ej,"makani_number")))
        if ej.get("landlord_name"):     kv.append(_kv("Landlord Name", _v(ej,"landlord_name")))
        if ej.get("landlord_nationality"): kv.append(_kv("Landlord Nationality", _v(ej,"landlord_nationality")))
        if ej.get("property_manager"): kv.append(_kv("Property Manager", _v(ej,"property_manager")))
        panels.append({"title": "EJARI — Tenancy Contract", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 5 — MOA Details
    # ══════════════════════════════════════════════════════════════════════════
    if moa_up:
        kv = []
        if moa.get("contract_number"):    kv.append(_kv("Contract No.", _v(moa, "contract_number")))
        if moa.get("moa_date"):           kv.append(_kv("Date of MOA", _fmt_date(moa["moa_date"])))
        if moa.get("company_name"):       kv.append(_kv("Company Name (English)", _v(moa, "company_name")))
        if moa.get("company_name_arabic"):kv.append(_kv("Company Name (Arabic)", _v(moa, "company_name_arabic")))
        if moa.get("legal_form"):         kv.append(_kv("Legal Form", _v(moa, "legal_form")))
        if moa.get("company_duration"):   kv.append(_kv("Company Duration", _v(moa, "company_duration")))
        if moa.get("financial_year"):     kv.append(_kv("Financial Year", _v(moa, "financial_year")))
        if moa.get("disputes_jurisdiction"): kv.append(_kv("Disputes Jurisdiction", _v(moa, "disputes_jurisdiction")))
        if moa.get("share_capital"):      kv.append(_kv("Share Capital", _v(moa, "share_capital")))
        if moa.get("shares_count"):       kv.append(_kv("Number of Shares", _v(moa, "shares_count")))
        if moa.get("capital_currency"):   kv.append(_kv("Currency", _v(moa, "capital_currency")))
        if moa.get("capital_deposited"):  kv.append(_kv("Capital Deposited", _v(moa, "capital_deposited")))
        if moa.get("statutory_reserve"):  kv.append(_kv("Statutory Reserve", _v(moa, "statutory_reserve")))
        if moa.get("owner_name"):         kv.append(_kv("Owner / Shareholder", _v(moa, "owner_name")))
        if moa.get("owner_nationality"):  kv.append(_kv("Owner Nationality", _v(moa, "owner_nationality")))
        if moa.get("owner_shares"):       kv.append(_kv("Shareholding Detail", _v(moa, "owner_shares")))
        if moa.get("owner_liability"):    kv.append(_kv("Owner Liability", _v(moa, "owner_liability")))
        if moa.get("owner_residence"):    kv.append(_kv("Owner Residence", _v(moa, "owner_residence")))
        if moa.get("manager_name"):       kv.append(_kv("Manager Name", _v(moa, "manager_name")))
        if moa.get("manager_nationality"):kv.append(_kv("Manager Nationality", _v(moa, "manager_nationality")))
        if moa.get("manager_role"):       kv.append(_kv("Manager Role", _v(moa, "manager_role")))
        if moa.get("manager_appointment_term"): kv.append(_kv("Appointment Term", _v(moa, "manager_appointment_term")))
        if moa.get("manager_residence"):  kv.append(_kv("Manager Residence", _v(moa, "manager_residence")))
        if moa.get("signing_authority"):  kv.append(_kv("Signing Authority", _v(moa, "signing_authority")))
        if moa.get("authorised_signatory"):kv.append(_kv("Authorised Signatory", _v(moa, "authorised_signatory")))
        if moa.get("signing_mode"):       kv.append(_kv("Signing Mode", _v(moa, "signing_mode")))
        if moa.get("bank_open_close"):    kv.append(_kv("Open / Close Bank Accounts", _v(moa, "bank_open_close")))
        if moa.get("bank_operate"):       kv.append(_kv("Operate Bank Accounts", _v(moa, "bank_operate")))
        if moa.get("bank_cheques"):       kv.append(_kv("Sign Cheques", _v(moa, "bank_cheques")))
        if moa.get("bank_transfer"):      kv.append(_kv("Transfer / Withdraw Funds", _v(moa, "bank_transfer")))
        if moa.get("bank_tenders"):       kv.append(_kv("Sign Tenders & Contracts", _v(moa, "bank_tenders")))
        if moa.get("bank_lc"):            kv.append(_kv("Issue Letters of Credit", _v(moa, "bank_lc")))
        if moa.get("bank_vat"):           kv.append(_kv("VAT / FTA Returns", _v(moa, "bank_vat")))
        if moa.get("bank_delegate"):      kv.append(_kv("Delegate Authority", _v(moa, "bank_delegate")))
        if kv: panels.append({"title": "Memorandum of Association (MOA)", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 6 — Insurance Details
    # ══════════════════════════════════════════════════════════════════════════
    if ins_up:
        kv = []
        if ins.get("insurer"):           kv.append(_kv("Insurer", _v(ins, "insurer")))
        if ins.get("insurer_arabic"):    kv.append(_kv("Insurer (Arabic)", _v(ins, "insurer_arabic")))
        if ins.get("policy_number"):     kv.append(_kv("Policy No.", _v(ins, "policy_number")))
        if ins.get("insured_name"):      kv.append(_kv("Insured Name", _v(ins, "insured_name")))
        if ins.get("insured_name_arabic"):kv.append(_kv("Insured Name (Arabic)", _v(ins, "insured_name_arabic")))
        if ins.get("coverage_type"):     kv.append(_kv("Coverage Type", _v(ins, "coverage_type")))
        if ins.get("sum_insured"):       kv.append(_kv("Sum Insured", _v(ins, "sum_insured")))
        if ins.get("premium"):           kv.append(_kv("Premium", _v(ins, "premium")))
        if ins.get("deductible"):        kv.append(_kv("Deductible / Excess", _v(ins, "deductible")))
        if ins.get("valid_from"):        kv.append(_kv("Valid From", _fmt_date(ins["valid_from"])))
        if ins.get("valid_to"):          kv.append(_kv("Valid To", ins_lbl, ins_sym))
        if kv: panels.append({"title": "Insurance Certificate", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 7 — VAT Certificate Details
    # ══════════════════════════════════════════════════════════════════════════
    if vat_up:
        kv = []
        if vat.get("trn"):                   kv.append(_kv("Tax Registration No. (TRN)", _v(vat, "trn")))
        if vat.get("company_name"):          kv.append(_kv("Registered Name", _v(vat, "company_name")))
        if vat.get("company_name_arabic"):   kv.append(_kv("Registered Name (Arabic)", _v(vat, "company_name_arabic")))
        if vat.get("effective_date"):        kv.append(_kv("Effective Date", _fmt_date(vat["effective_date"])))
        if vat.get("registered_address"):    kv.append(_kv("Registered Address", _v(vat, "registered_address")))
        if vat.get("registered_address_arabic"): kv.append(_kv("Address (Arabic)", _v(vat, "registered_address_arabic")))
        if vat.get("return_period"):         kv.append(_kv("Return Period", _v(vat, "return_period")))
        if vat.get("registration_type"):     kv.append(_kv("Registration Type", _v(vat, "registration_type")))
        kv.append(_kv("Expiry", "No expiry date — ongoing registration"))
        if kv: panels.append({"title": "VAT Registration Certificate", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 8 — Personal Documents
    # ══════════════════════════════════════════════════════════════════════════
    if has_multi_partner:
        for pidx, pdoc in enumerate(partner_docs_list, 1):
            pname = pdoc.get("partner_name", f"Partner {pidx}")
            kv = []
            for dtype, dkey, fields in [
                ("EMIRATES ID", "emirates_id", [
                    ("Name", "holder_name", False), ("ID No.", "id_number", False),
                    ("Date of Birth", "date_of_birth", True), ("Nationality", "nationality", False)]),
                ("PASSPORT", "passport", [
                    ("Name", "holder_name", False), ("Passport No.", "passport_number", False),
                    ("Nationality", "nationality", False), ("Date of Birth", "date_of_birth", True),
                    ("Place of Birth", "place_of_birth", False), ("Issue Date", "issue_date", True)]),
                ("UAE RESIDENCE VISA", "residence_visa", [
                    ("Name", "holder_name", False), ("Visa / Permit No.", "visa_number", False),
                    ("File No.", "file_number", False), ("Profession", "profession", False),
                    ("Sponsor / Employer", "employer", False)]),
            ]:
                d = dict(pdoc.get(dkey) or {})
                d.pop("error", None)
                if not d:
                    continue
                kv.append(_kv(f"─── {dtype} ──────────────────────────────────────────", ""))
                for label, key, is_date in fields:
                    if d.get(key):
                        kv.append(_kv(label, _fmt_date(d[key]) if is_date else _v(d, key)))
                if d.get("expiry_date"):
                    d_lbl, d_sym = _expiry_label(d["expiry_date"], today)
                    kv.append(_kv("Expiry Date", d_lbl, d_sym))
            if kv:
                panels.append({"title": f"Personal Documents — {pname}", "type": "kv", "rows": kv})
    elif pp_up or eid_up or visa_up:
        kv = []
        if eid_up and eid:
            kv.append(_kv("─── EMIRATES ID ─────────────────────────────────────────────────", ""))
            if eid.get("holder_name"):    kv.append(_kv("Name",          _v(eid,"holder_name")))
            if eid.get("id_number"):      kv.append(_kv("ID No.",         _v(eid,"id_number")))
            if eid.get("date_of_birth"):  kv.append(_kv("Date of Birth",  _fmt_date(eid["date_of_birth"])))
            if eid.get("nationality"):    kv.append(_kv("Nationality",    _v(eid,"nationality")))
            if eid.get("expiry_date"):    kv.append(_kv("Expiry Date",    eid_lbl, eid_sym))
        if pp_up and pp:
            kv.append(_kv("─── PASSPORT ────────────────────────────────────────────────────", ""))
            if pp.get("holder_name"):     kv.append(_kv("Name",           _v(pp,"holder_name")))
            if pp.get("passport_number"): kv.append(_kv("Passport No.",   _v(pp,"passport_number")))
            if pp.get("nationality"):     kv.append(_kv("Nationality",    _v(pp,"nationality")))
            if pp.get("date_of_birth"):   kv.append(_kv("Date of Birth",  _fmt_date(pp["date_of_birth"])))
            if pp.get("place_of_birth"):  kv.append(_kv("Place of Birth", _v(pp,"place_of_birth")))
            if pp.get("issue_date"):      kv.append(_kv("Issue Date",     _fmt_date(pp["issue_date"])))
            if pp.get("expiry_date"):     kv.append(_kv("Expiry Date",    pp_lbl, pp_sym))
        if visa_up and visa:
            kv.append(_kv("─── UAE RESIDENCE VISA ──────────────────────────────────────────", ""))
            if visa.get("holder_name"):   kv.append(_kv("Name",                _v(visa,"holder_name")))
            if visa.get("visa_number"):   kv.append(_kv("Visa / Permit No.",   _v(visa,"visa_number")))
            if visa.get("file_number"):   kv.append(_kv("File No.",            _v(visa,"file_number")))
            if visa.get("uid_number"):    kv.append(_kv("Unified No. (UID)",   _v(visa,"uid_number")))
            if visa.get("profession"):    kv.append(_kv("Profession",          _v(visa,"profession")))
            if visa.get("employer"):      kv.append(_kv("Sponsor / Employer",  _v(visa,"employer")))
            if visa.get("nationality"):   kv.append(_kv("Nationality",         _v(visa,"nationality")))
            if visa.get("gender"):        kv.append(_kv("Gender",              _v(visa,"gender")))
            if visa.get("date_of_birth"): kv.append(_kv("Date of Birth",       _fmt_date(visa["date_of_birth"])))
            if visa.get("passport_number"): kv.append(_kv("Passport No.",      _v(visa,"passport_number")))
            if visa.get("issue_date"):    kv.append(_kv("Issue Date",          _fmt_date(visa["issue_date"])))
            if visa.get("expiry_date"):   kv.append(_kv("Expiry Date",         visa_lbl, visa_sym))
        if kv: panels.append({"title": "Personal Documents", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 6 — Address Cross-Verification
    # ══════════════════════════════════════════════════════════════════════════
    if (tl_up and ej_up) or (tl_up and vat_up) or (ej_up and vat_up):
        tl_unit  = _v(tl,  "unit_number", "");     ej_unit  = _v(ej, "unit_number", "")
        tl_bldg  = _v(tl,  "building_name", "");   ej_bldg  = _v(ej, "building_name", "")
        tl_area  = _v(tl,  "area", "");            ej_area  = _v(ej, "area", "")
        tl_parc  = _v(tl,  "parcel_id", "");       ej_parc  = _v(ej, "land_dm_parcel_id", "")
        tl_addr  = _v(tl,  "registered_address", ""); vat_addr = _v(vat, "registered_address", "")
        ej_tenant  = _v(ej, "tenant_name", "");    tl_company = _v(tl, "company_name", "")
        ej_lic_no  = _v(ej, "licence_number", ""); tl_lic_no  = _v(tl, "license_number", "")

        if vat_up and (tl_up or ej_up):
            hdrs = ["Field", "Trade Licence", "EJARI", "VAT Certificate", "Match"]
            mr_rows = [
                _mr("Company / Tenant Name",
                    [tl_company or "—", ej_tenant or "—", _v(vat,"company_name") or "—"],
                    _msym(_match3(tl_company, ej_tenant, _v(vat,"company_name","")))),
                _mr("Full Address",
                    [tl_addr or "—", (ej_bldg+" "+ej_area).strip() or "—", vat_addr or "—"],
                    _msym(_match3(tl_addr, (ej_bldg+" "+ej_area).strip(), vat_addr))),
                _mr("Unit No.",      [tl_unit or "—", ej_unit or "—", "—"], _msym(_match2(tl_unit, ej_unit))),
                _mr("Building",      [tl_bldg or "—", ej_bldg or "—", "—"], _msym(_match2(tl_bldg, ej_bldg))),
                _mr("Area",          [tl_area or "—", ej_area or "—", "—"], _msym(_match2(tl_area, ej_area))),
                _mr("Parcel ID / Land DM", [tl_parc or "—", ej_parc or "—", "—"], _msym(_match2(tl_parc, ej_parc))),
                _mr("Licence No.",   [tl_lic_no or "—", ej_lic_no or "—", "—"], _msym(_match2(tl_lic_no, ej_lic_no))),
            ]
        else:
            hdrs = ["Field", "Trade Licence", "EJARI", "Match"]
            mr_rows = [
                _mr("Company / Tenant Name", [tl_company or "—", ej_tenant or "—"],  _msym(_match2(tl_company, ej_tenant))),
                _mr("Licence No.",           [tl_lic_no  or "—", ej_lic_no  or "—"], _msym(_match2(tl_lic_no,  ej_lic_no))),
                _mr("Unit No.",              [tl_unit    or "—", ej_unit    or "—"], _msym(_match2(tl_unit,    ej_unit))),
                _mr("Building",              [tl_bldg    or "—", ej_bldg    or "—"], _msym(_match2(tl_bldg,    ej_bldg))),
                _mr("Area",                  [tl_area    or "—", ej_area    or "—"], _msym(_match2(tl_area,    ej_area))),
                _mr("Parcel ID / Land DM",   [tl_parc    or "—", ej_parc    or "—"], _msym(_match2(tl_parc,    ej_parc))),
            ]
        panels.append({"title": "Address Cross-Verification", "type": "match", "headers": hdrs, "rows": mr_rows})
        for r in mr_rows:
            if r["match"] == "warn":
                _flag(f"Address Mismatch — {r['field']}", " vs ".join(hdrs[1:-1]),
                      r["field"], r["values"][0] if r["values"] else "—",
                      r["values"][1] if len(r["values"]) > 1 else "—",
                      "Verify with client and update the relevant document.")

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 7 — Name Verification: Trade Licence vs MOA
    # ══════════════════════════════════════════════════════════════════════════
    if tl_up and moa_up:
        moa_shareholders_v = moa.get("shareholders") if isinstance(moa.get("shareholders"), list) else []
        moa_shareholders_v = [s for s in moa_shareholders_v if isinstance(s, dict) and s.get("name")]
        moa_managers_v = moa.get("managers") if isinstance(moa.get("managers"), list) else []
        moa_managers_v = [m for m in moa_managers_v if isinstance(m, dict) and m.get("name")]

        def _join_v(items, key):
            return ", ".join(_s(x.get(key, "")).strip() for x in items
                             if isinstance(x, dict) and x.get(key)) or ""

        tl_own  = _v(tl,  "owner_name", "")
        moa_own = _join_v(moa_shareholders_v, "name") or _v(moa, "owner_name", "")
        tl_onat = _v(tl,  "owner_nationality", "")
        moa_onat = _s((moa_shareholders_v[0].get("nationality") if moa_shareholders_v else "")
                     or moa.get("owner_nationality", "")).strip()
        tl_oshr = _v(tl,  "owner_share", "")
        moa_oshr = _s((moa_shareholders_v[0].get("shares") or moa_shareholders_v[0].get("share_percentage")
                       if moa_shareholders_v else "") or moa.get("owner_shares", "")).strip()
        tl_mgr  = _v(tl,  "manager_name", "")
        moa_mgr = _join_v(moa_managers_v, "name") or _v(moa, "manager_name", "")
        tl_mnat = _v(tl,  "manager_nationality", "")
        moa_mnat = _s((moa_managers_v[0].get("nationality") if moa_managers_v else "")
                     or moa.get("manager_nationality", "")).strip()
        tl_mrol = _v(tl,  "manager_role", "")
        moa_mrol = _s((moa_managers_v[0].get("role") if moa_managers_v else "")
                     or moa.get("manager_role", "") or moa.get("signing_authority", "")).strip()

        def _multi_name_match_p(a: str, b: str) -> str:
            if not a or a == "—" or not b or b == "—":
                return "—"
            a_parts = [p.strip() for p in str(a).split(",") if p.strip()]
            b_parts = [p.strip() for p in str(b).split(",") if p.strip()]
            if len(a_parts) <= 1 and len(b_parts) <= 1:
                return _match2(a, b)
            small, large = (a_parts, b_parts) if len(a_parts) <= len(b_parts) else (b_parts, a_parts)
            for sname in small:
                if not any(_names_match(sname, lname) for lname in large):
                    return "⚠"
            return "✓"

        nv_rows = [
            _mr("Owner Name",          [tl_own  or "—", moa_own  or "—"], _msym(_multi_name_match_p(tl_own,  moa_own))),
            _mr("Owner Nationality",   [tl_onat or "—", moa_onat or "—"], _msym(_match2(tl_onat, moa_onat))),
            _mr("Owner Share",         [tl_oshr or "—", moa_oshr or "—"], _msym(_percent_match(tl_oshr, moa_oshr))),
            _mr("Manager Name",        [tl_mgr  or "—", moa_mgr  or "—"], _msym(_multi_name_match_p(tl_mgr,  moa_mgr))),
            _mr("Manager Nationality", [tl_mnat or "—", moa_mnat or "—"], _msym(_match2(tl_mnat, moa_mnat))),
            _mr("Manager Role",        [tl_mrol or "—", moa_mrol or "—"], _msym(_match2(tl_mrol, moa_mrol))),
        ]
        panels.append({"title": "Name Verification — Trade Licence vs MOA", "type": "match",
                        "headers": ["Field", "Trade Licence", "MOA", "Match"], "rows": nv_rows})
        _name_field_actions_p = {
            "Owner Name":          ("Name Mismatch",        "Align owner name spelling across Trade Licence and MOA."),
            "Owner Nationality":   ("Nationality Mismatch", "Verify owner nationality on Trade Licence and MOA."),
            "Owner Share":         ("Share Mismatch",       "Verify shareholding percentage on Trade Licence and MOA."),
            "Manager Name":        ("Name Mismatch",        "Align manager name spelling across Trade Licence and MOA."),
            "Manager Nationality": ("Nationality Mismatch", "Verify manager nationality on Trade Licence and MOA."),
            "Manager Role":        ("Role Mismatch",        "Verify manager role on Trade Licence and MOA."),
        }
        for r in nv_rows:
            if r["match"] == "warn":
                ftype, action = _name_field_actions_p.get(
                    r["field"], ("Mismatch", "Verify and align values across Trade Licence and MOA."))
                _flag(f"{ftype} — {r['field']}", "Trade Licence vs MOA",
                      r["field"], r["values"][0], r["values"][1], action)

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL 8 — Personal Document Cross-Verification
    # ══════════════════════════════════════════════════════════════════════════
    if has_multi_partner:
        co_name = _v(tl, "company_name", "") or _v(moa, "company_name", "")
        for pidx, pdoc in enumerate(partner_docs_list, 1):
            pname = pdoc.get("partner_name", f"Partner {pidx}")
            ppp   = dict(pdoc.get("passport") or {});       ppp.pop("error", None)
            peid  = dict(pdoc.get("emirates_id") or {});     peid.pop("error", None)
            pvisa = dict(pdoc.get("residence_visa") or {});  pvisa.pop("error", None)
            p_eid_name  = _v(peid,  "holder_name", "") if peid  else ""
            p_pp_name   = _v(ppp,   "holder_name", "") if ppp   else ""
            p_visa_name = _v(pvisa, "holder_name", "") if pvisa else ""
            p_present = [x for x in [p_eid_name, p_pp_name, p_visa_name] if x and x != "—"]

            if len(p_present) >= 2:
                m3 = _match3(p_eid_name, p_pp_name, p_visa_name)
                panels.append({
                    "title": f"Name Cross-Match — {pname}",
                    "type": "match",
                    "headers": ["Field", "Emirates ID", "Passport", "UAE Residence Visa", "Match"],
                    "rows": [_mr("Full Name",
                                 [p_eid_name or "—", p_pp_name or "—", p_visa_name or "—"],
                                 _msym(m3))]
                })
                if m3 == "⚠":
                    _flag(f"Name Mismatch — {pname}", "EID / Passport / Visa",
                          "Holder Name",
                          f"EID: {p_eid_name} | PP: {p_pp_name} | Visa: {p_visa_name}", "",
                          "Names must be consistent across all personal identity documents.")

            if ppp.get("passport_number") and pvisa.get("passport_number"):
                pp_no   = _v(ppp,   "passport_number")
                visa_pp = _v(pvisa, "passport_number")
                mpp = _match2(pp_no, visa_pp)
                panels.append({
                    "title": f"Passport Number — {pname}",
                    "type": "match",
                    "headers": ["Field", "Passport", "UAE Residence Visa", "Match"],
                    "rows": [_mr("Passport No.", [pp_no, visa_pp], _msym(mpp))]
                })
                if mpp == "⚠":
                    _flag(f"Passport No. Mismatch — {pname}", "Passport vs Visa",
                          "Passport No.", pp_no, visa_pp, "Verify passport number.")

            if p_present and pname and pname != "—":
                corp_rows_p: list[dict] = []
                for lbl, nm in [("Emirates ID", p_eid_name), ("Passport", p_pp_name),
                                ("Residence Visa", p_visa_name)]:
                    if nm and nm != "—":
                        m = _match2(nm, pname)
                        corp_rows_p.append(_mr(f"{lbl} vs Partners Annex", [nm, pname], _msym(m)))
                        if m == "⚠":
                            _flag(f"Name Mismatch — {lbl} ({pname})", f"{lbl} vs Partners Annex",
                                  "Holder Name", nm, pname, "Verify name spelling.")
                if corp_rows_p:
                    panels.append({
                        "title": f"Name vs Partners Annex — {pname}",
                        "type": "match",
                        "headers": ["Comparison", "Personal Document", "Partners Annex Name", "Match"],
                        "rows": corp_rows_p
                    })

            if pvisa.get("employer") and co_name and co_name != "—":
                emp = _v(pvisa, "employer")
                m = _match2(emp, co_name)
                panels.append({
                    "title": f"Employer vs Company — {pname}",
                    "type": "match",
                    "headers": ["Field", "Visa — Employer / Sponsor", "Company Name (TL / MOA)", "Match"],
                    "rows": [_mr("Employer / Company", [emp, co_name], _msym(m))]
                })
                if m == "⚠":
                    _flag(f"Employer Mismatch — {pname}", "Visa vs TL/MOA",
                          "Employer / Sponsor", emp, co_name,
                          "Verify sponsor on visa matches company name.")

    elif eid_up or pp_up or visa_up:
        eid_name  = _v(eid,  "holder_name", "") if eid_up  else ""
        pp_name   = _v(pp,   "holder_name", "") if pp_up   else ""
        visa_name = _v(visa, "holder_name", "") if visa_up else ""
        present   = [x for x in [eid_name, pp_name, visa_name] if x and x != "—"]

        if len(present) >= 2:
            m3 = _match3(eid_name, pp_name, visa_name)
            panels.append({
                "title": "Personal Document Name Cross-Match",
                "type": "match",
                "headers": ["Field", "Emirates ID", "Passport", "UAE Residence Visa", "Match"],
                "rows": [_mr("Full Name",
                             [eid_name or "—", pp_name or "—", visa_name or "—"],
                             _msym(m3))]
            })
            if m3 == "⚠":
                _flag("Personal Name Mismatch", "EID / Passport / Visa", "Holder Name",
                      f"EID: {eid_name} | PP: {pp_name} | Visa: {visa_name}", "",
                      "Names must be consistent across all personal identity documents.")

        if pp_up and visa_up and pp.get("passport_number") and visa.get("passport_number"):
            pp_no   = _v(pp,   "passport_number")
            visa_pp = _v(visa, "passport_number")
            mpp = _match2(pp_no, visa_pp)
            panels.append({
                "title": "Passport Number — Passport vs Residence Visa",
                "type": "match",
                "headers": ["Field", "Passport", "UAE Residence Visa", "Match"],
                "rows": [_mr("Passport No.", [pp_no, visa_pp], _msym(mpp))]
            })
            if mpp == "⚠":
                _flag("Passport No. Mismatch", "Passport vs Residence Visa",
                      "Passport No.", pp_no, visa_pp,
                      "Verify — visa may reference an expired passport.")

        ref_name = (_v(moa,"owner_name","") or _v(tl,"owner_name","") or
                    _v(moa,"manager_name","") or _v(tl,"manager_name",""))
        if ref_name and ref_name != "—" and present:
            corp_rows: list[dict] = []
            for label, name in [("Emirates ID", eid_name), ("Passport", pp_name), ("Residence Visa", visa_name)]:
                if name and name != "—":
                    m = _match2(name, ref_name)
                    corp_rows.append(_mr(f"{label} vs TL / MOA", [name, ref_name], _msym(m)))
                    if m == "⚠":
                        _flag(f"Name Mismatch — {label} vs Corporate", f"{label} vs TL/MOA",
                              "Holder Name", name, ref_name,
                              "Verify name spelling across personal and corporate documents.")
            if corp_rows:
                panels.append({
                    "title": "Personal Name vs Corporate Documents",
                    "type": "match",
                    "headers": ["Comparison", "Personal Document", "TL / MOA Name", "Match"],
                    "rows": corp_rows
                })

        if visa_up and visa.get("employer"):
            co  = _v(tl,"company_name","") or _v(moa,"company_name","")
            emp = _v(visa, "employer")
            if co and co != "—":
                m = _match2(emp, co)
                panels.append({
                    "title": "Employer (Visa) vs Company Name",
                    "type": "match",
                    "headers": ["Field", "Visa — Employer / Sponsor", "Company Name (TL / MOA)", "Match"],
                    "rows": [_mr("Employer / Company", [emp, co], _msym(m))]
                })
                if m == "⚠":
                    _flag("Employer Mismatch", "Residence Visa vs TL/MOA",
                          "Employer / Sponsor", emp, co,
                          "Verify sponsor on visa matches the company on Trade Licence.")

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL — Board Resolution Status
    # ══════════════════════════════════════════════════════════════════════════
    if moa_up or br_up:
        kv = []
        has_bank_open   = bool(moa.get("bank_open_close"))
        has_bank_cheque = bool(moa.get("bank_cheques"))
        has_bank_xfer   = bool(moa.get("bank_transfer"))
        moa_sufficient  = has_bank_open or has_bank_cheque or has_bank_xfer

        if moa_up:
            kv.append(_kv("Bank Account Opening",
                         "✓ Authorised — per MOA" if has_bank_open else "✗ Not stated in MOA",
                         "✓" if has_bank_open else "✗"))
            kv.append(_kv("Cheque Signing",
                         "✓ Authorised — per MOA" if has_bank_cheque else "✗ Not stated in MOA",
                         "✓" if has_bank_cheque else "✗"))
            kv.append(_kv("Fund Transfer",
                         "✓ Authorised — per MOA" if has_bank_xfer else "✗ Not stated in MOA",
                         "✓" if has_bank_xfer else "✗"))
            kv.append(_kv("Board Resolution Required",
                         "✗ No — MOA sufficient" if moa_sufficient
                         else "✓ Yes — MOA does not grant banking authority",
                         "✓" if moa_sufficient else "⚠"))

        if br_up:
            kv.append(_kv("─── BOARD RESOLUTION ────────────────────────────────────────────", ""))
            kv.append(_kv("Resolution Type", _v(br, "resolution_type")))
            kv.append(_kv("Resolution Date", _fmt_date(br.get("resolution_date"))))
            kv.append(_kv("Named Signatory", _v(br, "signatory_name")))
            kv.append(_kv("Named Bank(s)", _v(br, "named_banks")))
            kv.append(_kv("Notarised", "✓ Yes" if br.get("notarised") else "✗ No / Not confirmed",
                         "✓" if br.get("notarised") else "⚠"))

            if not moa_sufficient and not br_up:
                _flag("Banking Authority Missing", "MOA", "Banking Authority",
                      "MOA does not explicitly grant banking authority.", "",
                      "Provide notarised Board Resolution.")
        elif not moa_sufficient and moa_up:
            _flag("Board Resolution Required", "MOA", "Banking Authority",
                  "MOA silent on banking powers — Board Resolution not provided.", "",
                  "Provide notarised Board Resolution authorising signatory for banking.")

        panels.append({"title": "Board Resolution Status", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL — Physical Presence & POA Status
    # ══════════════════════════════════════════════════════════════════════════
    if eid_up or visa_up or pp_up or poa_up:
        kv = []
        signatory_name = (
            _v(moa, "authorised_signatory", "") or
            _v(moa, "manager_name", "") or _v(tl, "manager_name", "") or
            _v(br, "signatory_name", "") or
            _v(moa, "owner_name", "") or _v(tl, "owner_name", "") or
            _v(pp, "holder_name", "") or _v(eid, "holder_name", "")
        )
        kv.append(_kv("Authorised Signatory", signatory_name or "—"))

        eid_ok = eid_up and eid_sym == "✓"
        pp_ok  = pp_up  and pp_sym  == "✓"
        visa_ok = visa_up and visa_sym == "✓"

        kv.append(_kv("EID Valid", "✓ Valid" if eid_ok else ("✗ Expired/Missing" if eid_up else "— Not uploaded"),
                      "✓" if eid_ok else ("✗" if eid_up else "")))
        kv.append(_kv("Passport Valid", "✓ Valid" if pp_ok else ("✗ Expired/Missing" if pp_up else "— Not uploaded"),
                      "✓" if pp_ok else ("✗" if pp_up else "")))
        kv.append(_kv("UAE Visa Valid", "✓ Valid" if visa_ok else ("✗ Expired/Missing" if visa_up else "— Not uploaded"),
                      "✓" if visa_ok else ("✗" if visa_up else "")))
        kv.append(_kv("Can Attend Bank",
                      "✓ Yes" if (eid_ok and pp_ok and visa_ok) else "⚠ Blocked — see flags",
                      "✓" if (eid_ok and pp_ok and visa_ok) else "⚠"))

        if poa_up:
            kv.append(_kv("─── POWER OF ATTORNEY ──────────────────────────────────────────", ""))
            kv.append(_kv("Grantor", _v(poa, "grantor_name")))
            kv.append(_kv("Grantee (Attorney)", _v(poa, "grantee_name")))
            if poa.get("scope_description"):
                kv.append(_kv("Scope", _v(poa, "scope_description")))
            if poa.get("named_banks"):
                kv.append(_kv("Named Bank(s)", _v(poa, "named_banks")))
            kv.append(_kv("Notarised", "✓ Yes" if poa.get("notarised") else "✗ No / Not confirmed",
                         "✓" if poa.get("notarised") else "⚠"))
            if poa.get("expiry_date"):
                poa_lbl, poa_sym = _expiry_label(poa.get("expiry_date"), today)
                kv.append(_kv("Expiry", poa_lbl, poa_sym))

        panels.append({"title": "Physical Presence & POA Status", "type": "kv", "rows": kv})

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL — Partners / Shareholders
    # ══════════════════════════════════════════════════════════════════════════
    if pa_up:
        kv = []
        if pa.get("company_name"):
            kv.append(_kv("Company Name", _v(pa, "company_name")))
        if pa.get("licence_number"):
            kv.append(_kv("Licence No.", _v(pa, "licence_number")))
        partners = pa.get("partners", [])
        if isinstance(partners, list):
            for idx_p, partner in enumerate(partners, 1):
                if not isinstance(partner, dict):
                    continue
                kv.append(_kv(f"─── PARTNER {idx_p} ─────────────────────────────────────────────", ""))
                kv.append(_kv("Name", _v(partner, "name")))
                kv.append(_kv("Nationality", _v(partner, "nationality")))
                kv.append(_kv("Shareholding", _v(partner, "share_percentage")))
                is_corp = partner.get("is_corporate", False)
                kv.append(_kv("Type", "CORPORATE ENTITY" if is_corp else "Natural Person",
                             "⚠" if is_corp else "✓"))
                if is_corp:
                    _flag("Corporate Shareholder Identified", "Partners Annex",
                          "Shareholder Type",
                          f"{_v(partner, 'name')} holds {_v(partner, 'share_percentage')} — enhanced KYC required.",
                          "",
                          "Full corporate KYC required: Certificate of Incorporation, MOA/AOA, "
                          "Register of Shareholders/Directors, Good Standing, Board Resolution, "
                          "UBO passports. All foreign docs need 4-stage attestation chain.")
        panels.append({"title": "Partners / Shareholders — Annex", "type": "kv", "rows": kv})

    out = {"company_name": company_name, "panels": panels, "flags": flags}
    if analysis is not None:
        out["analysis"] = analysis
    return out
